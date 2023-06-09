from docx import Document
import docx

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_SECTION

from .classes import Pasta
from . import funcoes_gerais


#document = Document()
#styles =  document.styles

#Style Paragraph
#p = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
#p.font.name = "Calibri"
#p.font.size = Pt(11)

#Style Heading 2222
'''
h2 = styles.add_style("H2", WD_STYLE_TYPE.PARAGRAPH)
h2.base_style = styles["Heading 2"]
h2.font.name = "Calibri"
h2.font.size = Pt(13)
h2.font.color.rgb = RGBColor(78, 129, 189)
h2.font.bold = False

#Style Heading 3
h3 = styles.add_style("H3", WD_STYLE_TYPE.PARAGRAPH)
h3.base_style = styles["Heading 3"]
h3.font.name = "Calibri"
h3.font.size = Pt(12)
h3.font.color.rgb = RGBColor(78, 129, 189)
h3.font.bold = False
'''
def fun_peticao_02(p_pasta,p_autor,p_nr_processo,p_comarca,p_uf,p_cliente):
    doc = docx.Document()
    doc.add_heading('Peticao do documento', 0)
    para = doc.add_paragraph(
    '''GeeksforGeeks is a Computer Science portal for geeks.''')

    # Adding more content to paragraph and applying underline to them
    para.add_run(
    ''' It contains well written, well thought and well-explained ''').underline = True

    # Adding more content to paragraph
    para.add_run('''computer science and programming articles, quizzes etc.''')

    # Now save the document to a location

    doc.save('/home/ubuntu/documentos/teste2.docx')
    return 'teste2.docx'



def fun_peticao_01(pasta_saj):
    document = docx.Document()
    sec1 =  document.sections[-1]
    ft1 = sec1.footer
    hd1 =  sec1.header

    ft1_pg =  ft1.add_paragraph('João Barbosa Advogado Associados')
    ft1_pg.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hd1_pg =  hd1.add_paragraph(pasta_saj.cod_cliente+' C3/'+pasta_saj.pasta+'/'+pasta_saj.cobertura)
    hd1_pg.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    styles = document.styles
    p = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    p.font.name = "Times New Roman"
    p.font.size = Pt(13)
    #p.font.color.rgb=RGBColor(79, 129, 189)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #p.paragraph_format.first_line_indent = Inches(0.5)
    #p.paragraph_format.line_spacing = Inches(0.35)


    p2 = styles.add_style("Paragraph-2", WD_STYLE_TYPE.PARAGRAPH)
    p2.font.name = "Times New Roman"
    p2.font.size = Pt(12)
    #p2.font.color.rgb=RGBColor(79, 129, 189)
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #p2.paragraph_format.first_line_indent = Inches(0.5)
    p2.paragraph_format.left_indent = Inches(1.5)

    p3 = styles.add_style("Paragraph-3", WD_STYLE_TYPE.PARAGRAPH)
    p3.font.name = "Times New Roman"
    p3.font.size = Pt(12)
    p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p4 = styles.add_style("Paragraph-4", WD_STYLE_TYPE.PARAGRAPH)
    p4.font.name = "Times New Roman"
    p4.font.size = Pt(13)
    p4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p4.font.bold = True

    p41 = styles.add_style("Paragraph-41", WD_STYLE_TYPE.PARAGRAPH)
    p41.font.name = "Times New Roman"
    p41.font.size = Pt(12)
    p41.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p41.font.bold = True
    p41.font.italic = True


    p5 = styles.add_style("Paragraph-5", WD_STYLE_TYPE.PARAGRAPH)
    p5.font.name = "Times New Roman"
    p5.font.size = Pt(12)
    p5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #p5.paragraph_format.line_spacing = Inches(1)

    p5.paragraph_format.space_before = Pt(5)
    p5.paragraph_format.space_after = Pt(0.5)

    p6 = styles.add_style("Paragraph-6", WD_STYLE_TYPE.PARAGRAPH)
    p6.font.name = "Times New Roman"
    p6.font.size = Pt(12)
    p6.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #p5.paragraph_format.line_spacing = Inches(0.1)

    p6.paragraph_format.space_before = Pt(0.5)
    p6.paragraph_format.space_after = Pt(5)


    #document.add_paragraph(pasta_saj.cod_cliente+' C3/'+pasta_saj.pasta+'/'+pasta_saj.cobertura)
    #document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    document.add_picture('/home/fernandopaz/projetos/imagens/logo_jbaa.png', width=Inches(2.0)) 
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    #document.add_picture('/home/ubuntu/documentos/logo_jbaa.png', width=Inches(2.0))


    document.add_paragraph('EXMO SR. DR. JUIZ DE DIREITO DO(A) '+(pasta_saj.juizo).upper()+' DA COMARCA DE '+pasta_saj.comarca+'/'+pasta_saj.uf, style='Paragraph-4')

    document.add_paragraph('',style='Paragraph-4')
    document.add_paragraph('',style='Paragraph-4')

    document.add_paragraph('Processo n. '+pasta_saj.nr_processo, style='Paragraph-41')



    #para = document.add_paragraph(pasta_saj.cliente+', previamente qualificada nos autos do processo em epígrafe, neste ato, representada por seus advogados que esta subscrevem, nos autos da ', style='Paragraph')

    para = document.add_paragraph('',style='Paragraph')
    para.add_run(pasta_saj.cliente).bold = True
    para.add_run(', previamente qualificada nos autos do processo em epígrafe, neste ato, representada por seus advogados que esta subscrevem, nos autos da ')
    para.add_run('AÇÃO DE COBRANÇA DE SEGURO DPVAT.').bold = True
    para.add_run(' que lhe promove ')
    para.add_run(pasta_saj.autor).bold = True
    para.add_run(', em trâmite perante este Douto Juízo, vem respeitosamente, à presença de V. Exa., ')
    para.add_run('requerer o DESARQUIVAMENTO, a fim de viabilizar a DEVOLUÇÃO DOS HONORÁRIOS PERICIAIS PAGOS EM DUPLICIDADE (depósito judicial e ofício único de pagamento.').bold=True
    #document.add_page_break()
    '''
    under_and_bold = para.add_run('AÇÃO DE COBRANÇA DE SEGURO DPVAT')
    under_and_bold.underline=True
    under_and_bold.bold=True
    para.add_run(', que lhe promove ')
    para.add_run(pasta_saj.autor).bold=True
    para.add_run(', em trâmite perante este Douto Juízo, vem respeitosamente, à presença de V. Exa., requerer o DESARQUIVAMENTO, a fim de viabilizar a DEVOLUÇÃO DOS HONORÁRIOS PERICIAIS PAGOS EM DUPLICIDADE (depósito judicial e ofício único de pagamento.')
    '''
    document.add_paragraph('Consoante se verifica nos autos e da documentação que segue em anexo, houve depósito a título de pagamento de honorários periciais, em cumprimento à intimação de fls., contudo, o processo foi relacionado para evento de mutirão de perícias, ocasião em que houve o pagamento da prova através de ofício único, restando, portanto, pagamento em duplicidade.',style='Paragraph')

    para = document.add_paragraph('Desta forma, com fulcro no art. 906, parágrafo único do CPC, requer a Ré que Vossa Excelência se digne determinar a expedição de ',style='Paragraph')
    run = para.add_run('OFÍCIO DE TRANSFERÊNCIA DIRETA no montante do valor depositado, ')
    run.font.bold =  True
    run.font.underline = True
    para.add_run('com seus acréscimos legais, em favor da ')
    para.add_run('SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A., CNPJ/MF: 09.248.608/0001-04, ').bold = True
    para.add_run('autorizando ao Banco depositante a efetuar transferência direta na ')
    para.add_run('conta corrente nº 644000-2, Agência: 1912-7, BANCO DO BRASIL S.A').bold = True
    #document.add_page_break()


    sec2= document.add_section(WD_SECTION.NEW_PAGE)
    hd2 = sec2.header
    ft2 = sec2.footer

    hd2.is_linked_to_previous = False
    ft2.is_linked_to_previous = False


    para = document.add_paragraph('Necessário esclarecer que a expedição da ordem de pagamento deverá ser nominal à ',style = 'Paragraph' )
    run = para.add_run('SEGURADORA LÍDER DOS CONSÓRCIOS DO SEGURO DPVAT S/A,')
    run.font.bold = True
    run.font.underline = True
    para.add_run(' pois foi a empresa que custeou com o depósito como também é a gestora dos ')
    para.add_run('Consórcios do Seguro DPVAT nos termos do art. 5º, §3º, da Resolução CNSP de nº 154,').bold = True
    para.add_run(' sendo a única e exclusiva beneficiária de reembolso da quantia disponível ao juízo.')


    para = document.add_paragraph('Reforçando o acima exposto, temos que as regras e os critérios para o DPVAT referentes aos sinistros ocorridos ', style='Paragraph')
    para.add_run('até 31 de dezembro de 2020').bold = True
    para.add_run(' estão estabelecidas, também, na Resolução n.º 399 do CNSP de 29/12/2020.')

    document.add_paragraph('A referida Resolução prevê, no seu artigo 21, a competência da Seguradora Líder:',style='Paragraph')

    para = document.add_paragraph('Art. 21. ',style = 'Paragraph-2')
    para.add_run('A seguradora líder').bold = True
    para.add_run(' do Consórcio DPVAT será ')
    para.add_run('responsável').bold = True
    para.add_run(' pela gestão e operacionalização do seguro ')
    para.add_run('DPVAT').bold = True
    para.add_run(' referentes, exclusivamente, ')
    para.add_run('aos sinistros ocorridos até 31 de dezembro de 2020').bold = True
    para.add_run(' (run-off), inclusive em relação às respectivas ações judiciais posteriormente ajuizadas.')

    document.add_paragraph('Vejamos, agora, o art. 1º da Resolução 400 do CNSP de 29/12/2020:', style='Paragraph')

    para = document.add_paragraph('Art. 1º ',style = 'Paragraph-2')
    para.add_run('Ratificar que a Seguradora Líder').bold = True 
    para.add_run(' do Consórcio do Seguro DPVAT S.A. será a ')
    para.add_run('responsável').bold = True
    para.add_run(' pela gestão e operacionalização do seguro ')
    para.add_run('DPVAT').bold = True
    para.add_run(' referentes, exclusivamente, ')
    para.add_run('aos sinistros ocorridos até 31 de dezembro de 2020,').bold = True
    para.add_run(' inclusive em relação às respectivas ações judiciais posteriormente ajuizadas.')

    document.add_paragraph('Requer ainda, seja determinado que o banco depositante junte aos autos o respectivo comprovante da transferência realizada através de TED da quantia expedida mediante oficio, possibilitando ao patrono da Ré realizar prestação de contas com maior clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.',style='Paragraph')

    document.add_paragraph('Por fim, que seja observado exclusivamente o nome do advogado '+pasta_saj.conv_nome+', '+ pasta_saj.conv_oab+ ' para efeito de intimações futuras, sob pena de nulidade das mesmas.',style='Paragraph')

    document.add_paragraph('Termos em que,',style='Paragraph-3')


    document.add_paragraph('Pede Juntada.',style='Paragraph-3')
    document.add_paragraph(funcoes_gerais.local_e_data(pasta_saj.comarca),style='Paragraph-3')

    document.add_paragraph('João Barbosa',style='Paragraph-5')

    document.add_paragraph(pasta_saj.oabjb,style='Paragraph-6')


    document.add_paragraph(pasta_saj.conv_nome,style='Paragraph-5')
    document.add_paragraph('OAB '+pasta_saj.conv_oab,style='Paragraph-6')

    nome_do_arquivo='Peticao_Desarquiv_Devolucao_HonDuplicidade_'+pasta_saj.cod_cliente+'.docx'

    document.save('/home/fernandopaz/projetos/documentos/'+nome_do_arquivo)
    return nome_do_arquivo

