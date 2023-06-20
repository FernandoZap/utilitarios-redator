# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Length
import docx

import os
import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches,Cm
from . import funcoes_gerais,estilos,funcoes_banco,funcoes_impressao

document = Document()


def peticoes(teses,dados_compl,dados_pasta):
    #dados_view = funcoes_banco.f001_sql(pasta,1)
    pasta = dados_pasta['pasta']
    cod_cliente = dados_pasta['cod_cliente']



    doc = docx.Document()

    styles = doc.styles
    p = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    p.font.name = "Calibri"
    p.font.size = Pt(11)
    p.font.bold = True
    #p.font.color.rgb=RGBColor(79, 129, 189)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #p.paragraph_format.first_line_indent = Inches(0.5)


    p2 = styles.add_style("Paragraph-2", WD_STYLE_TYPE.PARAGRAPH)
    p2.font.name = "Calibri"
    p2.font.size = Pt(11)
    #p2.font.color.rgb=RGBColor(79, 129, 189)
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #p2.paragraph_format.first_line_indent = Inches(0.5)
    #p2.paragraph_format.left_indent = Inches(0.5)


    # config inicial / logo / rodape
    funcoes_impressao.pagina_inicial(doc,pasta)
    funcoes_impressao.sinteseDosFatos(doc)

    for tese, valor in teses.items():
        if valor=='S':
            if tese=='Nulidade':
                t001_teseNulidade(dados_pasta,dados_compl,pasta,doc)
            if tese=='IntimacaoDoMP':
                t002_teseIntimacaoDoMP(doc)
            if tese=='CoisaJulgada':
                t003_teseDaCoisaJulgada(dados_pasta,dados_compl,doc)
            if tese=='Litispendencia':
                t004_teseDaLitispendencia(dados_pasta,dados_compl,doc)
            if tese=='OmissaoPrescricao':
                t005_teseOmissaoPrescricao(dados_pasta,dados_compl,doc)
            if tese=='OmissaoPagamentoAdm':
                t006_teseOmissaoPagamento(dados_pasta,dados_compl,doc)
            if tese=='OmissaoInadimplente':
                t007_teseOmissaoInadimplente(dados_pasta,dados_compl,doc)
            if tese=='OmissaoLesaoPreExistente':
                t008_teseOmissaoLesaoPreExistente(dados_pasta,dados_compl,doc)
            if tese=='OmissaoRegulacao8':
                t009_teseOmissaoRegulacao8(dados_pasta,dados_compl,doc)
            if tese=='OmissaoConsectariosLegais':
                t010_teseOmissaoConsectariosLegais(dados_pasta,dados_compl,doc)



    # assinatura
    funcoes_impressao.pagina_conclusao(doc,pasta)
    funcoes_impressao.pagina_de_encerramento(doc,pasta)

    nome=cod_cliente+'_'+pasta+'_Embargos_Omissao' + funcoes_gerais.data_doc() +'.docx'
    nome_documento = os.environ.get('DIR_DOCUMENTOS')+'/'+nome
    doc.save(nome_documento)

    return nome


def t001_teseNulidade(dados_pasta,dados_compl,pasta,doc):

    publicando_nome = dados_pasta['publicando_nome']
    publicando_oab =  'OAB: '+dados_pasta['publicando_oab']
    publicando_nome_oab = dados_pasta['publicando_nome']+' - OAB: '+dados_pasta['publicando_oab']

    doc.add_paragraph('DA TEMPESTIVIDADE',style='Paragraph')

    #doc.add_paragraph('NULIDADE DE INFORMAÇÃO',style='Paragraph')



    para = doc.add_paragraph('Inicialmente, cumpre observar que foi publicado dia ',style='Paragraph-2')
    para.add_run(dados_compl['data_public']).bold =  True
    para.add_run(', no Diário da Justiça Eletrônico, a r. decisão exarada, como se verifica na colação abaixo:' )

    paragraph = doc.add_paragraph('',style='Paragraph')
    run2 = paragraph.add_run('"COLAR A PUBLICACÃO"')
    blue = RGBColor(79, 129, 189)
    run2.font.color.rgb =  blue


    doc.add_paragraph('Desta feita, a Seguradora permanecia no aguardo da devida publicação para que pudesse verificar a '+
    'intenção em recorrer, e ofertar sua peça tempestivamente, o que o faz sob ancorada no princípio de '+
    'celeridade e economia processual.',style='Paragraph-2')

    paragraph = doc.add_paragraph('Como se vê não foram respeitadas as exigências de Publicidade dos atos praticados, tendo em vista que foi,', style='Paragraph-2')
    paragraph.add_run(' requerido na peça de bloqueio (fls.), que futuras publicações fossem feitas em nome do patrono da Apelante ')
    paragraph.add_run(publicando_nome).bold = True


    para = doc.add_paragraph('Conclui-se, portanto, que em nenhum momento o ',style='Paragraph-2')
    para.add_run('r. decisum esteve à disposição da Seguradora para ciência')
    para.add_run(' e eventual manifestação nos autos.')

    doc.add_paragraph('Afinal não é possível que a Seguradora, com seu grandioso número de causas, possua o controle e tenha a '+
    'possibilidade de organizar suas publicações com seus números de processo.',style='Paragraph-2')

    doc.add_paragraph('Por tal motivo, inclusive, é que se indica os nomes dos patronos a saírem a publicação realizada, eis que se '+
    'torna uma forma mais fácil de proceder o acompanhamento processual.',style='Paragraph-2')

    para = doc.add_paragraph('Assim, repita-se, ',style='Paragraph-2')
    runner = para.add_run('NÃO HOUVE PUBLICAÇÃO DA D. SENTENÇA,')
    runner.bold = True
    runner.underline = True
    para.add_run(' o que ocasionou a perda do prazo para manifestação nos autos.').bold = True

    doc.add_paragraph('Neste sentido, os requisitos formais para a validade do ato de comunicação processual, fundamental para '+
    'a aplicação dos regimes de preclusão e desenvolvimento dos atos processuais, não atendeu aos critérios '+
    'formais de sua realização.',style='Paragraph-2')

    para = doc.add_paragraph('Conclui-se, portanto, que ',style='Paragraph-2')
    runner =  para.add_run('em nenhum momento o r. decisum esteve à disposição da Recorrente para ciência,')
    runner.bold = True
    para.add_run(' haja vista que NÃO foi publicada em nome do patrono constituído nos autos.')

    para = doc.add_paragraph('Assim, requer o recebimento da presente peça processual, ante as nulidades suscitadas. ',style='Paragraph-2')

    return doc

def teseIntimacaoDoMP_bak(dados_pasta,dados_compl,pasta,doc):


    publicando_nome = dados_pasta['publicando_nome']
    publicando_oab =  'OAB: '+dados_pasta['publicando_oab']
    elementos = []

    doc.add_paragraph('DA SÍNTESE DOS FATOS E DA OMISSÃO', style='Paragraph')

    doc.add_paragraph('Com a mais a respeitosa vênia, assim o fazendo, '+
    'afigura-se a v. decisão omissa em pontos essenciais, justificando o cabimento dos '+
    'presentes Embargos de Declaração, a fim de que essa V. Exa. decida-os e confira os '+
    'efeitos integrativos ao respeitável decisum.',style='Paragraph-2')


    doc.add_paragraph('Frisa-se que na d. sentença exarada, verifica-se grave OMISSÃO, que devem ser supridas ou sanadas por meio dos presentes embargos, sendo certo que o recurso não objetiva rediscutir a matéria, mas afastar os vícios constatados no julgado.',style='Paragraph-2')

    doc.add_paragraph('Com todo o respeito, a Embargante informa que nos casos em que há interesses de incapazes sendo discutido em determinada demanda, deve ser intimado o Ministério Público, órgão fiscalizador da Lei, para que se pronuncie sobre a necessidade de sua intervenção.',style='Paragraph-2')

    doc.add_paragraph('Cumpre informar, no caso dos autos, o autor é menor, e figura como autor na presente demanda, figurando como representante, seu genitor, contudo, em que pese tenha haja o pedido de intimação do MP na peça de bloqueio, não se observa menção a este respeito na sentença prolatada.',style='Paragraph-2')

    doc.add_paragraph('Urge ressaltar, a necessidade da prática deste ato, de intimação do MP, não por uma faculdade, mas um comando imposto pelo Código de Processo Civil, que traz inclusive, quando ausente tal intimação, uma possibilidade do reconhecimento de uma nulidade.',style='Paragraph-2')

    doc.add_paragraph('Ante o exposto e da patente necessidade de intimação do Ministério Público para fins de atender ao disposto nos artigos 178, II c/c 279 do CPC, requer seja verificada a omissão informada e a consequente intimação do Parquet para acompanhar o feito.',style='Paragraph-2')

    doc.add_paragraph('CONCLUSÃO', style='Paragraph')

    doc.add_paragraph('São essas as razões pelas quais a embargante confia, espera e requer sejam acolhidos e providos os presentes Embargos Declaratórios, enfrentado o ponto OMISSO, conferido efeitos integrativos para o fim de prover integralmente, tudo por ser medida de direito e irretorquível JUSTIÇA!',style='Paragraph-2')

    return doc

def t002_teseIntimacaoDoMP(doc):


    doc.add_paragraph('DA NECESSÁRIA DE INTIMAÇÃO DO MP', style='Paragraph')

    doc.add_paragraph('Com todo respeito, a Embargante informa que nos casos em que há interesses de incapazes sendo discutido em determinada demanda, deve ser intimado o Ministério Público, órgão fiscalizador '+
    'da Lei, para que se pronuncie sobre a necessidade de sua intervenção.')

    doc.add_paragraph('Cumpre informar, nos casos dos autos, o autor é menor, e figura como autor na presente demanda, figurando como representante, seu genitor, contudo, em que pese tenha '+
    'haja o pedido de intimação do MP na peça de bloqueio, não se observa menção a este respeito '+
    'na sentença prolatada.')

    doc.add_paragraph('Urge ressaltar, a necessidade da prática deste ato, de intimação do MP, não por uma faculdade, mas um comando imposto pelo Código do Processo Civil, '+
    'que traz inclusive, quando ausente tal intimação, uma possibilidade do reconhecimento deu uma nulidade.')

    doc.add_paragraph('Ante o exposto e da patente necessidade de intimação do Ministério Público para fins de atender ao disposto nos artigos 178, II c/c 279 do CPC, requer seja verificada a '+
    'omissão informada e a consequente intimação do Parquet para acompanhar o feito.')


def t003_teseDaCoisaJulgada(dados_pasta,dados_compl,doc):


    nr_processo=dados_pasta['nr_processo']
    nr_processo_vinculado=dados_compl['num_processo_vinculado']
    juizo_vinculado=dados_compl['juizo_vinculado']
    juizo=dados_pasta['juizo'] + ' DE ' + dados_pasta['comarca'] + ' - ' + dados_pasta['estado']

    doc.add_paragraph('DA COISA JULGADA',style='Paragraph')

    doc.add_paragraph('Preliminarmente, informa da existência de outra demanda idêntica a presente, ou seja, com as mesmas partes, pedido e causa de pedir, a qual fora registrada '+
    'sob o número '+ nr_processo_vinculado  +', e tramitou perante o Juízo da '+ juizo_vinculado  +', tendo havido trânsito em julgado de decisão de mérito, '+
    'fazendo-se coisa julgada material, conforme comprovam as cópias inclusas.')
    last_pg=doc.paragraphs[-1]
    last_pg.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph

    doc.add_paragraph('Desta feita, manifesta a tríplice identidade entre a presente demanda e aquela supramencionada, pelo que se requer o acolhimento desta preliminar, '+
    'a fim de se julgar EXTINTO o feito, nos termos do art. 485, V, do CPC.')
    last_pg=doc.paragraphs[-1]
    last_pg.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph

    doc.add_paragraph('Por fim, pugna-se pela condenação da parte a todos os consectários legais, inclusive custas processuais, honorários advocatícios e ainda, '+
    'a condenação pela comprovada litigância de má-fé conforme disposto no artigo 77 da Lei Processual Civil.')

    last_pg=doc.paragraphs[-1]
    last_pg.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph



def t004_teseDaLitispendencia(dados_pasta,dados_compl,doc):


    num_processo_vinculado=dados_compl['num_processo_vinculado']
    juizo_vinculado=dados_compl['juizo_vinculado']

    doc.add_paragraph('DA LITISPENDENCIA', style='Paragraph')

    paragraph = doc.add_paragraph('Preliminarmente, informa da existência de outra demanda idêntica a presente, ou seja, '+
    'com as mesmas partes, pedido e causa de pedir, a qual fora registrada sob o número ')
    paragraph.add_run(num_processo_vinculado).bold =  True
    paragraph.add_run(', e tramita perante o Juízo da ' + juizo_vinculado  + ', conforme comprovam as cópias inclusas.')

    last_pg=doc.paragraphs[-1]
    last_pg.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph

    doc.add_paragraph('Como se sabe, o instituto da litispendência é matéria de ordem pública, razão pela qual '+
    'pode/deve ser alegada em qualquer tempo e grau de jurisdição, inclusive ser declarada ex officio pelo Juízo.')
    last_pg=doc.paragraphs[-1]
    last_pg.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph


    doc.add_paragraph('Desta feita, manifesta a tríplice identidade entre a presente demanda e aquela supramencionada, '+
    'pelo que se requer o acolhimento desta, a fim de se julgar EXTINTO o feito, sem resolução de mérito, nos termos do art. 485, V, do CPC.')
    last_pg=doc.paragraphs[-1]
    last_pg.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph

def t005_teseOmissaoPrescricao(dados_pasta,dados_compl,doc):

    doc.add_paragraph('DA PRESCRICAO', style='Paragraph')

    doc.add_paragraph('Com a mais a respeitosa vênia, na decisão proferida V. Exa. não se manifestou, '+
    'expressamente, sobre pontos importantes levantados na contestação, a respeito dos quais, deveria '+
    'ter-se pronunciado, justificando o cabimento dos presentes Embargos de Declaração, para que lhes '+
    'confira os efeitos integrativos ao respeitável decisum.')
    last_pg=doc.paragraphs[-1]
    last_pg.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph


    doc.add_paragraph('Conforme sustentado pela Embargante em sua peça de bloqueio o direito postulatório '+
    'está IRREMEDIAVELMENTE PRESCRITO.')
    last_pg=doc.paragraphs[-1]
    last_pg.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph


    doc.add_paragraph('Conforme amplamente demonstrado nos autos, trata-se, da chamada “prescrição extintiva”, '+
    'donde se depreende que o não uso do direito no tempo previsto, acarreta sua perda.')
    last_pg=doc.paragraphs[-1]
    last_pg.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph


    doc.add_paragraph('Neste ponto a d. sentença não dedicou uma palavra sequer à esta questão amplamente invocada. '+
    'Quedando-se omisso a este respeito e merecendo reforma.')
    last_pg=doc.paragraphs[-1]
    last_pg.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph



def t006_teseOmissaoPagamento(dados_pasta,dados_compl,doc):


    num_processo_vinculado=dados_compl['num_processo_vinculado']
    juizo_vinculado=dados_compl['juizo_vinculado']


    doc.add_paragraph('DO INCONTROVERSO PAGAMENTO ADMINISTRATIVO',style='Paragraph')


    doc.add_paragraph('Ocorre que com relação ao sinistro alegado na presente demanda houve pagamento administrativo, '+
    'a Embargante, reitera que o pagamento foi realizado em favor do Embargado, conforme consta dos documentos '+
    'acostados – isto, após meticulosa análise da documentação apresentada foi liberado o valor da indenização na monta '+
    'de R$2.362,50 (Dois mil trezentos...), trazemos a colação o comprovante de pagamento, vejamos:')


    doc.add_paragraph('COLACIONAR O RECIBO DO PAGAMENTO ADM')


    doc.add_paragraph('Portanto, necessária a apreciação das provas trazidas ao processo pela ora Embargante, '+
    'uma vez que não foi considerado pelo juízo sentenciante que o pagamento administrativo ora noticiado.')

    doc.add_paragraph('Destaca-se que o seguro DPVAT é alvo de fraudes a todo instante! Não que seja o caso desses '+
    'autos, mas as evidencias se relevam como tentativa da requerente em receber valor além do estabelecido por lei, '+
    'ocultando o fato de já ter recebido a quantia de R$2.362,50 (Dois mil trezentos...) na via administrativa.')

    doc.add_paragraph('Ressalte-se que a Embargante não está se omitindo ou procrastinando na presente demanda, muito '+
    'pelo contrário, busca a veracidade dos fatos, para a perfeita aplicação da justiça.')

    doc.add_paragraph('De acordo com os documentos anexados pela Embargante, nota-se que o pagamento da indenização '+
    'ora pleiteada já foi objeto de análise e pagamento em sede administrativa.')

    doc.add_paragraph('Diante do exposto, requer sejam acolhidos e providos os presentes Embargos Declaratórios, '+
    'enfrentado os pontos omissos suscitados, conferido os efeitos integrativos, por via de consequência modificativos, '+
    'para o fim de prover integralmente, para que sobre eles se pronuncie esse Ilustre Julgador, tudo por ser medida de direito e justiça.')

    doc.add_paragraph('Outrossim, informa a embargante que pelo fato dos presentes Embargos terem efeitos infringentes, '+
    'requer que seja feita a devida intimação da parte Embargada, para que esta venha responder as presentes alegações, '+
    'a fim de evitar violação ao direito constitucional da ampla defesa e contraditório.')



def t007_teseOmissaoInadimplente(dados_pasta,dados_compl,doc):


    num_processo_vinculado=dados_compl['num_processo_vinculado']
    juizo_vinculado=dados_compl['juizo_vinculado']

    doc.add_paragraph('DA INAPLENCIA DA PARTE AUTORA COM O SEGURO DPVAT',style='Paragraph')

    doc.add_paragraph('Conforme sustentado pela Embargante em sua peça de bloqueio a parte Embargada quando do sinistro estava inadimplente com o Seguro DPVAT.')

    doc.add_paragraph('Conforme amplamente demonstrado, estando o pagamento do DPVAT em atraso, o veículo não é considerado licenciado, '+
    'o proprietário deixa de ter direito à cobertura em caso de acidente e, o proprietário é obrigado a ressarcir as indenizações '+
    'eventualmente pagas às vítimas do acidente.')

    doc.add_paragraph('Neste ponto a r. Decisão não dedicou uma palavra sequer à esta questão amplamente invocada nos autos. '+
    'Quedando-se omissa a este respeito e merecendo reforma.')


def t008_teseOmissaoLesaoPreExistente(dados_pasta,dados_compl,doc):


    num_proc_lpe = dados_compl['num_proc_lpe']
    juizo_lpe = dados_compl['juizo_lpe']
    data_lpe = dados_compl['data_lpe']
    desc_lpe = dados_compl['desc_lpe']

    doc.add_paragraph('DESCABIMENTO DE RENOVAÇÃO DE PLEITO INDENIZATÓRIO', style='Paragraph')

    doc.add_paragraph('LESÃO PREEXISTENTE', style='Paragraph')

    doc.add_paragraph('Inicialmente, deve-se sopesar o fato da parte Embargada ter pleiteado judicialmente verba '+
    'indenizatória DPVAT, cujo processo tramitou na ' + juizo_lpe +', sendo autuado sob o nº. ' + num_proc_lpe + ', em virtude de acidente '+
    'automobilístico ocorrido em '+ data_lpe)

    doc.add_paragraph('Frisa-se que a parte Embargada requereu o recebimento do Seguro Obrigatório DPVAT nos autos da '+
    'ação supracitada em decorrência de ' + desc_lpe  +  ', ou seja, o requerente sustenta seu pleito '+
    'indenizatório em lesão idêntica a que fora recebida anteriormente.')

    doc.add_paragraph('Constata-se que os documentos acostados aos autos comprovam que o acidente que ocasionou a debilidade '+
    'permanente foi anterior ao narrado na inicial, não havendo, portanto, nexo de causalidade entre o novo acidente e a '+
    'lesão apresentada pela parte autora.')

    doc.add_paragraph('Deste modo, é irrefragável que a presente lide tem o mesmo pedido de outra ação que teve o mérito julgado, '+
    'uma vez que a parte sequer comprova que houve agravamento da lesão em virtude de um suposto novo acidente automobilístico.')



def t009_teseOmissaoRegulacao8(dados_pasta,dados_compl,doc):


    local_diligencia = dados_compl['local_diligencia']

    doc.add_paragraph('DO CERCEAMENTO DO DIREITO À PRODUÇÃO DA PROVA') 

    doc.add_paragraph('Assim, o i. Magistrado permaneceu silente quanto os pedidos de diligências solicitados pela embargante, '+
    'qual seja, ' + local_diligencia)

    doc.add_paragraph('Neste sentido, em virtude da ausência de análise do argumento relativo a fatos relevantes para o deslinde da '+
    'causa, restaram violados os Princípios da Ampla Defesa e do Contraditório, tendo em vista que as alegações suscitadas quanto '+
    'as irregularidades ocorridas no processo administrativo não foram objeto de apreciação por este i. Juízo.')

    doc.add_paragraph('Vale destacar que o cerceamento do direito à produção da prova viola os direitos processuais da Embargante, '+
    'direitos instaurados no cerne da própria concepção do Estado de Direito Democrático e protegidos pela ordem jurídica.')

    doc.add_paragraph('A Embargante, faz lembrar ao atento juízo que o seguro DPVAT é alvo de milhares de fraudes em todo o Brasil, '+
    'não que seja o caso da presente demanda, sem contar que os argumentos da Embargante são de substancial importância para se '+
    'desvelar os fatos controvertidos.')

    doc.add_paragraph('Neste sentido requer seja sanada a omissão apontada e V. Exa. se digne a determinar ' + local_diligencia + ' a '+
    'fim de que sejam prestados os devidos esclarecimentos pelos responsáveis, sem prejuízo do colhimento do depoimento pessoal da parte embargada.')


def t010_teseOmissaoConsectariosLegais(dados_pasta,dados_compl,doc):

    doc.add_paragraph('DOS CONSECTÁRIOS LEGAIS')

    doc.add_paragraph('Com todo o respeito a Embargante, vem, informar que houve omissão quanto a atualização do valor '+
    'indenizatório, ou seja, a D.  não se manifestou sobre a data inicial para o compito dos juros.')

    doc.add_paragraph('É pacífico o entendimento firmado pelo Superior Tribunal de Justiça de que a aplicação de correção '+
    'monetária e juros de mora são matérias de ordem pública, podendo ser conhecidas de ofício e tampouco conduz à '+
    'interpretação de ocorrência de preclusão consumativa, porquanto tais institutos são meros consectários legais da condenação.')

    doc.add_paragraph('Em relação aos juros de mora, o Colendo Superior Tribunal de justiça editou a Súmula nº 426 pacificando a '+
    'incidência a partir da citação.')

    doc.add_paragraph('Em relação a correção monetária, o Colendo Superior Tribunal de justiça editou a Súmula nº 580 pacificando sua '+
    'incidência a partir do evento danoso.')

    doc.add_paragraph('Neste ponto, requer seja verificada a omissão informada, devendo-se esclarecer se o valor arbitrado será atualizado e '+
    'caso sim, que seja observado os ditames legais previstos para a matéria in foco.')



def fun_nr_processo(numproc,numprocSE,uf):
    nr_processo=numproc

    if uf=='SE':
        if numprocSE!='':
            nr_processo=numprocSE
    return nr_processo
