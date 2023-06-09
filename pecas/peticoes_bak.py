# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Length
import docx

import os
import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches,Cm
from . import funcoes_gerais,estilos,funcoes_banco,funcoes_impressao

document = Document()


def modelos_peticoes(modelo,dados_da_pasta,info):
    cod_cliente = dados_da_pasta.cod_cliente
    pasta = dados_da_pasta.pasta

    #PXaragraph-1--Style-11
    #PXaragraph-2--Style-2
    #PXaragraph-4--Style-4

    doc = docx.Document()

    styles = doc.styles
    p = styles.add_style("Style-1", WD_STYLE_TYPE.PARAGRAPH)
    p.font.name = "Calibri"
    p.font.size = Pt(11)
    p.font.bold = True
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    styles = doc.styles
    p = styles.add_style("Style-11", WD_STYLE_TYPE.PARAGRAPH)
    p.font.name = "Calibri"
    p.font.size = Pt(11)
    #p.font.bold = True
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #p.paragraph_format.first_line_indent = Inches(0.5)


    p2 = styles.add_style("Style-2", WD_STYLE_TYPE.PARAGRAPH)
    p2.font.name = "Calibri"
    p2.font.size = Pt(11)
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    p3 = styles.add_style("Style-3", WD_STYLE_TYPE.PARAGRAPH)
    p3.font.name = "Calibri"
    p3.font.size = Pt(11)
    p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p3.paragraph_format.left_indent = Inches(1.5)

    p4 = styles.add_style("Style-4", WD_STYLE_TYPE.PARAGRAPH)
    p4.font.name = "Calibri"
    p4.font.size = Pt(11)
    p4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p4.font.bold = True

    p5 = styles.add_style("Style-40", WD_STYLE_TYPE.PARAGRAPH)
    p5.font.name = "Calibri"
    p5.font.size = Pt(15)
    p5.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p5.font.bold = True



    # config inicial / logo / rodape
    #funcoes_impressao.pagina_inicial_peticoes_(doc,dados_da_pasta.pasta)
    #funcoes_impressao.sinteseDosFatos(doc)
    nome_arquivo = montarNomeArquivo(modelo,cod_cliente,pasta)
    if modelo=='desarquivamentoDevHonDupl':
        funcoes_impressao.pagina_inicial_peticoes_desarquivamento(doc,dados_da_pasta,1,info) ##
        pet_001_DesarquivamentoDevHonDupl(dados_da_pasta,doc)
    elif modelo=='dispComprovanteTransfComDesarquiv':
        funcoes_impressao.pagina_inicial_peticoes_desarquivamento(doc,dados_da_pasta,2,info) ##
        pet_002_DispComprovanteTransfComDesarquiv(dados_da_pasta,doc)
    elif modelo=='juntadaDePagamentoDaCondenacao':
        funcoes_impressao.pagina_inicial_peticoes_desarquivamento(doc,dados_da_pasta,8,info)
        pet_003_JuntadaDePagamentoDaCondenacao(dados_da_pasta,doc)
    elif modelo=='devolucaoHPRN1':
        funcoes_impressao.pagina_inicial_peticoes_desarquivamento(doc,dados_da_pasta,4,info) ##
        pet_004_DevolucaoHPRN1(dados_da_pasta,doc)
    elif modelo=='devolucaoHPRN2':
        funcoes_impressao.pagina_inicial_peticoes_desarquivamento(doc,dados_da_pasta,5,info) ##
        pet_005_DevolucaoHPRN2(dados_da_pasta,doc)
    elif modelo=='devolucaoHPRN3':
        funcoes_impressao.pagina_inicial_peticoes_desarquivamento(doc,dados_da_pasta,6,info) ##
        pet_006_DevolucaoHPRN3(dados_da_pasta,doc)
    elif modelo=='devolucaoHPRN4':
        funcoes_impressao.pagina_inicial_peticoes_desarquivamento(doc,dados_da_pasta,6,info) ##
        pet_007_DevolucaoHPRN4(dados_da_pasta,doc)
    elif modelo=='devolucaoHPRN5':
        funcoes_impressao.pagina_inicial_peticoes_desarquivamento(doc,dados_da_pasta,7,info) ##
        pet_008_DevolucaoHPRN5(dados_da_pasta,doc)
    elif modelo=='devolucaoHPRN6':
        funcoes_impressao.pagina_inicial_peticoes_desarquivamento(doc,dados_da_pasta,7,info) ##
        pet_009_DevolucaoHPRN6(dados_da_pasta,doc)
    elif modelo=='devolucaoHPRN7':
        funcoes_impressao.pagina_inicial_peticoes_desarquivamento(doc,dados_da_pasta,7,info) ##
        pet_010_DevolucaoHPRN7(dados_da_pasta,doc)
    elif modelo=='devolucaoHPImprocedencia':
        funcoes_impressao.pagina_inicial_peticoes_devolucao(doc,dados_da_pasta) ##
        pet_011_DevolucaoHPImprocedencia(dados_da_pasta,doc)
    elif modelo=='devolucaoHPImprocedenciaCDesarquiv':
        funcoes_impressao.pagina_inicial_peticoes_devolucao(doc,dados_da_pasta) ##
        pet_012_DevolucaoHPImprocedenciaCDesarquiv(dados_da_pasta,doc)
    elif modelo=='dispComprovanteTransf':
        funcoes_impressao.pagina_inicial_peticoes_devolucao(doc,dados_da_pasta) ##
        pet_013_DispComprovanteTransf(dados_da_pasta,doc)
    elif modelo=='reiterandoExpedicaoOficioDev':
        funcoes_impressao.pagina_inicial_peticoes_devolucao(doc,dados_da_pasta) ##
        pet_014_ReiterandoExpedicaoOficioDev(dados_da_pasta,doc)
    elif modelo=='devolucaoPernambuco':
        funcoes_impressao.pagina_inicial_peticoes_devolucao(doc,dados_da_pasta) ##
        pet_015_DevolucaoPernambuco(dados_da_pasta,doc)
    elif modelo=='devHPExtintoSemResolucao':
        funcoes_impressao.pagina_inicial_peticoes_devolucao(doc,dados_da_pasta) ##
        pet_016_DevHPExtintoSemResolucao(dados_da_pasta,doc)
    elif modelo=='devHPExtintoCDesarquivamento':
        funcoes_impressao.pagina_inicial_peticoes_devolucao(doc,dados_da_pasta) ##
        pet_017_DevHPExtintoCDesarquivamento(dados_da_pasta,doc)
    elif modelo=='juntadaDeCustasFinais':
        funcoes_impressao.pagina_inicial_peticoes_desarquivamento(doc,dados_da_pasta,9,info)
        pet_018_JuntadaDeCustasFinais(dados_da_pasta,doc)


    # assinatura
    #funcoes_impressao.pagina_conclusao(doc,dados_da_pasta.pasta)
    funcoes_impressao.pagina_de_encerramento_peticoes(doc,pasta)

    #nome_documento="c:/Avulsos/testes/embargos" +funcoes_gerais.data_doc() +".docx"
    #nome_documento="c:/projetos/utiljb/staticfiles/static/imagens/embargos" +funcoes_gerais.data_doc() +".docx"
    nome_documento =  "/home/fernandopaz/projetos/documentos/"+nome_arquivo+'_'+funcoes_gerais.data_doc() + ".docx"
    nome=nome_arquivo+'_'+ funcoes_gerais.data_doc() +".docx"
    doc.save(nome_documento)

    return nome


def pet_001_DesarquivamentoDevHonDupl(dados_da_pasta,doc):
    doc.add_paragraph("Consoante se verifica nos autos e da documentação que segue em anexo, houve depósito a título de pagamento de honorários periciais, "+
    "em cumprimento à intimação de fls., contudo, o processo foi relacionado para evento de mutirão de perícias, ocasião em que houve o pagamento da prova "+
    "através de ofício único, restando, portanto, pagamento em duplicidade.", style="Style-2")

    para = doc.add_paragraph('Desta forma, com fulcro no art. 906, parágrafo único do CPC, requer a Ré que Vossa Excelência se digne determinar a expedição de ',style='Style-2')
    para.add_run('OFÍCIO DE TRANSFERÊNCIA DIRETA no montante do valor depositado, ').bold = True
    para.add_run('com seus acréscimos legais, em favor da ')
    para.add_run('SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A., CNPJ/MF: 09.248.608/0001-04, ').bold = True
    para.add_run('autorizando ao Banco depositante a efetuar transferência direta na ')
    para.add_run('conta corrente nº 644000-2, Agência: 1912-7, BANCO DO BRASIL S.A').bold = True

    para = doc.add_paragraph('Necessário esclarecer que a expedição da ordem de pagamento deverá ser nominal à ',style = 'Style-2' )
    para.add_run('SEGURADORA LÍDER DOS CONSÓRCIOS DO SEGURO DPVAT S/A,').bold = True
    para.add_run(' pois foi a empresa que custeou com o depósito como também é a gestora dos ')
    para.add_run('Consórcios do Seguro DPVAT nos termos do art. 5º, §3º, da Resolução CNSP de nº 154,').bold = True
    para.add_run(' sendo a única e exclusiva beneficiária de reembolso da quantia disponível ao juízo.')

    para = doc.add_paragraph('Reforçando o acima exposto, temos que as regras e os critérios para o DPVAT referentes aos sinistros ocorridos ', style='Style-2')
    para.add_run('até 31 de dezembro de 2020').bold = True
    para.add_run(' estão estabelecidas, também, na Resolução n.º 399 do CNSP de 29/12/2020.')

    doc.add_paragraph('A referida Resolução prevê, no seu artigo 21, a competência da Seguradora Líder:',style='Style-2')

    artigoVinteUmResolucao399(doc)

    doc.add_paragraph('Vejamos, agora, o art. 1º da Resolução 400 do CNSP de 29/12/2020:', style='Style-2')

    artigoPrimeiroResolucao400(doc)

    doc.add_paragraph('Requer ainda, seja determinado que o banco depositante junte aos autos o respectivo comprovante da transferência realizada através '+
    'de TED da quantia expedida mediante oficio, possibilitando ao patrono da Ré realizar prestação de contas com maior clareza e transparência, informando o '+
    'saldo líquido e a data exata da transferência realizada.',style='Style-2')

    doc.add_paragraph('Por fim, que seja observado exclusivamente o nome do advogado '+dados_da_pasta.conv_nome+', '+ dados_da_pasta.conv_oab+ ' para efeito de intimações futuras, sob pena de nulidade das mesmas.',style='Style-2')

    doc.add_paragraph('Termos em que,\nPede Juntada,',style='Style-11')


def pet_011_DevolucaoHPImprocedencia(dados_da_pasta,doc):


    doc.add_paragraph('Em cumprimento à determinação desse d. juízo, a ré procedeu com o pagamento dos honorários periciais.',style='Style-2')
    doc.add_paragraph('Contudo, diante da ausência da parte autora à prova designada, imprescindível para análise do pedido reclamado, o processo foi julgado '+
    'improcedente, decisão esta que já transitou em julgado, merecendo o aludido valor depositado a título de honorários periciais, ser restituído à parte ré.',style='Style-2')

    doc.add_paragraph('Ante o exposto, requer que seja expedido OFÍCIO DE TRANSFERÊNCIA DIRETA, nos termos do parágrafo único, do art. 906, CPC, para fins de '+
    'devolução à ré do valor depositado nos autos, conforme anexo, e seus acréscimos legais, em favor da SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A., CNPJ/MF: 09.248.608/0001-04, '+
    'autorizando ao Banco depositante a efetuar transferência na conta corrente nº 644000-2, Agência: 1912-7, do BANCO DO BRASIL S/A.',style='Style-2')

    para = doc.add_paragraph('Necessário esclarecer que a expedição da ordem de pagamento deverá ser nominal à ', style='Style-2')
    run = para.add_run('SEGURADORA LÍDER DOS CONSÓRCIOS DO SEGURO DPVAT S/A')
    run.font.bold = True
    run.font.underline =  True
    para.add_run(', pois foi a empresa que custeou com o depósito como também é a ')
    run = para.add_run('GESTORA')
    run.font.bold = True
    run.font.underline = True
    para.add_run(' dos ')
    para.add_run('Consórcios do Seguro DPVAT nos termos do art. 5º, §3º, da Resolução CNSP de nº 154').bold = True
    para.add_run(', sendo a única e exclusiva beneficiária de reembolso da quantia disponível ao juízo.')

    para = doc.add_paragraph('Reforçando o acima exposto, temos que as regras e os critérios para o DPVAT referentes aos sinistros ocorridos ',style='Style-2')
    run = para.add_run('até 31 de dezembro de 2020')
    run.font.bold = True
    para.add_run(' estão estabelecidas, também, na Resolução n.º 399 do CNSP de 29/12/2020.')

    doc.add_paragraph('A referida Resolução prevê, no seu artigo 21, a competência da Seguradora Líder:',style='Style-2')

    artigoVinteUmResolucao399(doc)

    doc.add_paragraph('Vejamos, agora, o art. 1º da Resolução 400 do CNSP de 29/12/2020:',style='Style-2')

    artigoPrimeiroResolucao400(doc)

    doc.add_paragraph('Requer ainda, seja determinado que o banco depositante junte aos autos o respectivo comprovante da transferência realizada através de TED da quantia expedida '+
    'mediante oficio, possibilitando ao patrono da Ré realizar prestação de contas com maior clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.',style='Style-2')

    doc.add_paragraph('Nestes termos,\nPede Deferimento,',style='Style-11')
    doc.add_paragraph('Pede Deferimento,',style='Style-11')

    return document

def pet_013_DispComprovanteTransf(dados_da_pasta,doc):

    doc.add_paragraph('Conforme consta nos autos, existem valores a serem restituidos à ré tendo sido a ordem de transferência determinado por este d. Juízo.',style = 'Style-2')

    doc.add_paragraph('Ocorre que, ainda que expedido ofício ao gerente da instituição financeira depositante, para que fosse realizada transferência de valores '+
    'em favor da seguradora Ré, não houve resposta do mesmo, com apresentação nos autos do respectivo comprovante.', style = 'Style-2')

    doc.add_paragraph('Assim, vem a Ré requerer a V. Exa., seja determinado que o banco depositante junte aos autos o respectivo comprovante de transferência '+
    'realizada através de TED da quantia determinada em ofício, possibilitando ao patrono da Ré realizar prestação de contas com maior '+
    'clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.', style = 'Style-2')

    doc.add_paragraph('Ademais, pugna-se que na requisição conste prazo para cumprimento da ordem judicial, sob pena de crime de desobediência, a fim de '+
    'empregar plena efetividade e previsibilidade ao comando.', style = 'Style-2')

    doc.add_paragraph('Nestes termos,\nPede Deferimento,',style='Style-11')


def pet_012_DevolucaoHPImprocedenciaCDesarquiv(dados_da_pasta,doc):


    doc.add_paragraph('Em cumprimento à determinação desse d. juízo, a ré procedeu com o pagamento dos honorários periciais.',style='Style-2')

    doc.add_paragraph('Contudo, diante da ausência da parte autora à prova designada, imprescindível para análise do pedido reclamado, o processo foi julgado '+
    'improcedente, decisão esta que já transitou em julgado, merecendo o aludido valor depositado a título de honorários periciais, ser restituído à parte ré.',style='Style-2')

    doc.add_paragraph('Ante o exposto, requer que seja expedido OFÍCIO DE TRANSFERÊNCIA DIRETA, nos termos do parágrafo único, do art. 906, CPC, para fins de '+
    'devolução à ré do valor depositado nos autos, conforme anexo, e seus acréscimos legais, em favor da SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A., CNPJ/MF: 09.248.608/0001-04, '+
    'autorizando ao Banco depositante a efetuar transferência na conta corrente nº 644000-2, Agência: 1912-7, do BANCO DO BRASIL S/A.',style='Style-2')

    para = doc.add_paragraph('Necessário esclarecer que a expedição da ordem de pagamento deverá ser nominal à ', style='Style-2')
    run = para.add_run('SEGURADORA LÍDER DOS CONSÓRCIOS DO SEGURO DPVAT S/A')
    run.font.bold = True
    run.font.underline =  True
    para.add_run(', pois foi a empresa que custeou com o depósito como também é a ')
    run = para.add_run('GESTORA')
    run.font.bold = True
    run.font.underline = True
    para.add_run(' dos ')
    para.add_run('Consórcios do Seguro DPVAT nos termos do art. 5º, §3º, da Resolução CNSP de nº 154').bold = True
    para.add_run(', sendo a única e exclusiva beneficiária de reembolso da quantia disponível ao juízo.')

    para = doc.add_paragraph('Reforçando o acima exposto, temos que as regras e os critérios para o DPVAT referentes aos sinistros ocorridos ',style='Style-2')
    run = para.add_run('até 31 de dezembro de 2020')
    run.font.bold = True
    para.add_run(' estão estabelecidas, também, na Resolução n.º 399 do CNSP de 29/12/2020.')

    doc.add_paragraph('A referida Resolução prevê, no seu artigo 21, a competência da Seguradora Líder:',style='Style-2')

    artigoVinteUmResolucao399(doc)

    doc.add_paragraph('Vejamos, agora, o art. 1º da Resolução 400 do CNSP de 29/12/2020:',style='Style-2')

    artigoPrimeiroResolucao400(doc)

    doc.add_paragraph('Requer ainda, seja determinado que o banco depositante junte aos autos o respectivo comprovante da transferência realizada através de TED da quantia expedida '+
    'mediante oficio, possibilitando ao patrono da Ré realizar prestação de contas com maior clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.',style='Style-2')

    doc.add_paragraph('Nestes termos,\nPede Deferimento,',style='Style-11')
    doc.add_paragraph('Pede Deferimento,',style='Style-11')

    return document

'''
def peticao_DispComprovanteTransf(dados_da_pasta,doc):

    doc.add_paragraph('Conforme consta nos autos, existem valores a serem restituidos à ré tendo sido a ordem de transferência determinado por este d. Juízo.',style = 'Style-2')

    doc.add_paragraph('Ocorre que, ainda que expedido ofício ao gerente da instituição financeira depositante, para que fosse realizada transferência de valores\
    em favor da seguradora Ré, não houve resposta do mesmo, com apresentação nos autos do respectivo comprovante.', style = 'Style-2')

    doc.add_paragraph('Assim, vem a Ré requerer a V. Exa., seja determinado que o banco depositante junte aos autos o respectivo comprovante de transferência\
    realizada através de TED da quantia determinada em ofício, possibilitando ao patrono da Ré realizar prestação de contas com maior\
    clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.', style = 'Style-2')

    doc.add_paragraph('Ademais, pugna-se que na requisição conste prazo para cumprimento da ordem judicial, sob pena de crime de desobediência, a fim de\
    empregar plena efetividade e previsibilidade ao comando.', style = 'Style-2')

    doc.add_paragraph('Nestes termos,\nPede Deferimento,',style='Style-11')
'''


def pet_002_DispComprovanteTransfComDesarquiv(dados_da_pasta,doc):

    doc.add_paragraph('Conforme consta nos autos, existem valores a serem restituidos à ré tendo sido a ordem de transferência determinado por este d. Juízo.',style = 'Style-2')

    doc.add_paragraph('Ocorre que, ainda que expedido ofício ao gerente da instituição financeira depositante, para que fosse realizada transferência de valores '+
    'em favor da seguradora Ré, não houve resposta do mesmo, com apresentação nos autos do respectivo comprovante.', style = 'Style-2')

    doc.add_paragraph('Assim, vem a Ré requerer a V. Exa., seja determinado que o banco depositante junte aos autos o respectivo comprovante de transferência '+
    'realizada através de TED da quantia determinada em ofício, possibilitando ao patrono da Ré realizar prestação de contas com maior '+
    'clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.', style = 'Style-2')

    doc.add_paragraph('Ademais, pugna-se que na requisição conste prazo para cumprimento da ordem judicial, sob pena de crime de desobediência, a fim de '+
    'empregar plena efetividade e previsibilidade ao comando.', style = 'Style-2')

    doc.add_paragraph('Nestes termos,\nPede Deferimento,',style='Style-11')


def pet_014_ReiterandoExpedicaoOficioDev(dados_da_pasta,doc):

    doc.add_paragraph('Conforme já peticionado nos autos, existem valores a serem devolvidos ao Réu, conforme toda documentação já apresentada e novamente juntada.',style='Style-2')

    doc.add_paragraph('Destarte, renova-se o pedido de devolução dos valores, através da expedição de OFÍCIO DE TRANSFERÊNCIA DIRETA no montante do valor depositado, '+
    'com seus acréscimos legais, em favor da SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A., CNPJ/MF: 09.248.608/0001-04, autorizando ao Banco depositante a '+
    'efetuar transferência direta na conta corrente nº 644000-2, Agência: 1912-7, BANCO DO BRASIL S.A.', style='Style-2')

    para = doc.add_paragraph('Necessário esclarecer que a expedição da ordem de pagamento deverá ser nominal à ',style='Style-2')
    run = para.add_run('SEGURADORA LÍDER DOS CONSÓRCIOS DO SEGURO DPVAT S/A')
    run.font.bold = True
    para.add_run(', pois foi a empresa que custeou com o depósito como também é a ')
    run = para.add_run('GESTORA')
    run.font.bold = True
    para.add_run(' dos ')
    run = para.add_run('Consórcios do Seguro DPVAT nos termos do art. 5º, §3º, da Resolução CNSP de nº 154')
    run.font.bold = True
    para.add_run(', sendo a única e exclusiva beneficiária de reembolso da quantia disponível ao juízo.')

    para = doc.add_paragraph('Reforçando o acima exposto, temos que as regras e os critérios para o DPVAT referentes aos sinistros ocorridos ', style='Style-2')
    run = para.add_run('até 31 de dezembro de 2020')
    run.font.bold = True
    para.add_run(' estão estabelecidas, também, na Resolução n.º 399 do CNSP de 29/12/2020.')

    doc.add_paragraph('A referida Resolução prevê, no seu artigo 21, a competência da Seguradora Líder:', style='Style-2')

    artigoVinteUmResolucao399(doc)

    doc.add_paragraph('Vejamos, agora, o art. 1º da Resolução 400 do CNSP de 29/12/2020:', style='Style-2')

    artigoPrimeiroResolucao400(doc)

    doc.add_paragraph('Requer ainda, seja determinado que o banco depositante junte aos autos o respectivo comprovante da transferência '+
    'realizada através de TED da quantia expedida mediante oficio, possibilitando ao patrono da Ré realizar prestação '+
    'de contas com maior clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.',style='Style-2')

    doc.add_paragraph('Nestes termos,\nPede Deferimento,',style='Style-11')


def pet_015_DevolucaoPernambuco(dados_da_pasta,doc):

    doc.add_paragraph('Em cumprimento à determinação desse d. juízo, a ré procedeu com o pagamento dos honorários periciais.', style='Style-2')

    doc.add_paragraph('Contudo, diante da ausência da parte autora à prova designada, imprescindível para análise do pedido reclamado, '+
    'o processo foi julgado improcedente, decisão esta que já transitou em julgado, merecendo o aludido valor depositado a título '+
    'de honorários periciais, ser restituído à parte ré.', style='Style-2')

    doc.add_paragraph('Ante o exposto, e de acordo com o Ato nº 759/2022 da Presidência do Tribunal de Justiça de Pernambuco, '+
    'requer o levantamento do valor depositado e seus acréscimos legais, por meio do Malote e Digital em favor da '+
    'SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A., CNPJ/MF: 09.248.608/0001-04, '+
    'conta corrente nº 644000-2, Agência: 1912-7, do BANCO DO BRASIL S/A.',style = 'Style-4')

    para =doc.add_paragraph('Necessário esclarecer que a expedição da ordem de pagamento deverá ser nominal à ', style='Style-2')
    run = para.add_run('SEGURADORA LÍDER DOS CONSÓRCIOS DO SEGURO DPVAT S/A, ')
    run.font.bold = True
    para.add_run('pois foi a empresa que custeou com o depósito como também é a gestora dos ')
    run = para.add_run('Consórcios do Seguro DPVAT nos termos do art. 5º, §3º, da Resolução CNSP de nº 154, ')
    run.font.bold = True
    para.add_run('sendo a única e exclusiva beneficiária de reembolso da quantia disponível ao juízo.')

    para = doc.add_paragraph('Reforçando o acima exposto, temos que as regras e os critérios para o DPVAT referentes aos sinistros ocorridos ')
    run = para.add_run('até 31 de dezembro de 2020')
    run.font.bold = True
    para.add_run(' estão estabelecidas, também, na Resolução n.º 399 do CNSP de 29/12/2020.')

    doc.add_paragraph('A referida Resolução prevê, no seu artigo 21, a competência da Seguradora Líder:')

    artigoVinteUmResolucao399(doc)

    doc.add_paragraph('Vejamos, agora, o art. 1º da Resolução 400 do CNSP de 29/12/2020:',style = 'Style-2')

    artigoPrimeiroResolucao400(doc)

    para = doc.add_paragraph('Requer ainda, seja determinado que o banco depositante junte aos autos o respectivo comprovante '+
    'da  quantia expedida mediante Malote Digital, conforme ',style = 'Style-2')
    run = para.add_run('Ato nº 759/2022 da Presidência do Tribunal de Justiça de Pernambuco,')
    run.font.bold = True
    para.add_run(' possibilitando ao patrono da Ré realizar prestação de contas com maior clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.')


def pet_017_DevHPExtintoCDesarquivamento(dados_da_pasta,doc):
    pet016_DevHPExtintoSemResolucao(dados_da_pasta,doc)


def pet016_DevHPExtintoSemResolucao(dados_da_pasta,doc):

    doc.add_paragraph('Em cumprimento à determinação desse d. juízo, a ré procedeu com o pagamento dos honorários periciais.', style='Style-2')

    doc.add_paragraph('Contudo, diante da ausência da parte autora à prova designada, imprescindível para análise do pedido reclamado, '+
    'o processo foi julgado extinto sem resolução de mérito, decisão esta que já transitou em julgado, merecendo o aludido valor '+
    'depositado a título de honorários periciais, ser restituído à parte ré.', style='Style-2')

    doc.add_paragraph('Ante o exposto, requer que seja expedido OFÍCIO DE TRANSFERÊNCIA DIRETA, nos termos do parágrafo único, '+
    'do art. 906, CPC, para fins de devolução à ré do valor depositado nos autos, conforme anexo, e seus acréscimos legais, em '+
    'favor da SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A., CNPJ/MF: 09.248.608/0001-04, autorizando '+
    'ao Banco depositante a efetuar transferência na conta corrente nº 644000-2, Agência: 1912-7, do BANCO DO BRASIL S/A.', style='Style-2')


    para =doc.add_paragraph('Necessário esclarecer que a expedição da ordem de pagamento deverá ser nominal à', style='Style-2')
    run = para.add_run('SEGURADORA LÍDER DOS CONSÓRCIOS DO SEGURO DPVAT S/A,')
    run.font.bold = True
    para.add_run(' pois foi a empresa que custeou com o depósito como também é a ')
    run = para.add_run('GESTORA')
    run.font.bold = True
    para.add_run(' dos ')
    run = para.add_run('Consórcios do Seguro DPVAT nos termos do art. 5º, §3º, da Resolução CNSP de nº 154, ')
    run.font.bold = True
    para.add_run('sendo a única e exclusiva beneficiária de reembolso da quantia disponível ao juízo.')


    para = doc.add_paragraph('Reforçando o acima exposto, temos que as regras e os critérios para o DPVAT referentes aos sinistros ocorridos ', style = 'Style-2')
    run = para.add_run('até 31 de dezembro de 2020')
    run.font.bold = True
    para.add_run(' estão estabelecidas, também, na Resolução n.º 399 do CNSP de 29/12/2020.')

    doc.add_paragraph('A referida Resolução prevê, no seu artigo 21, a competência da Seguradora Líder:')

    artigoVinteUmResolucao399(doc)

    doc.add_paragraph('Vejamos, agora, o art. 1º da Resolução 400 do CNSP de 29/12/2020:',style = 'Style-2')

    artigoPrimeiroResolucao400(doc)

    doc.add_paragraph('Requer ainda, seja determinado que o banco depositante junte aos autos o respectivo comprovante da transferência '+
    'realizada através de TED da quantia expedida mediante oficio, possibilitando ao patrono da Ré realizar prestação de contas com maior '+
    'clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.', style = 'Style-2')

    doc.add_paragraph('Nestes termos,\nPede Deferimento,',style='Style-11')

def pet_003_JuntadaDePagamentoDaCondenacao(dados_da_pasta,doc):
    styles = doc.styles
    p = styles.add_style("Style-6", WD_STYLE_TYPE.PARAGRAPH)
    p.font.name = "Calibri"
    p.font.size = Pt(11)
    p.font.bold = True
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


    doc.add_paragraph('Assim, pugna a ré pela intimação da parte autora nos termos do art. 526, §1º, NCPC, havendo extinção com a concordância expressa ou em sendo ultrapassado '+
    'o prazo de 05 dias sem manifestação, deverá ser extinta a execução nos termos do art. 526, §3º c/c 924, II, NCPC.', style='Style-6')

    para = doc.add_paragraph('Por fim, que seja observado exclusivamente o nome ')
    if dados_da_pasta.publicando_sexo=='M':
        para.add_run('do advogado ')
    else:
        para.add_run('da advogada ')
    para.add_run(dados_da_pasta.publicando_nome+' '+dados_da_pasta.publicando_oab).bold = True
    para.add_run(', para efeito de intimações futuras, sob pena de nulidade das mesmas.')

    doc.add_paragraph('Termos em que,\nPede Juntada,',style='Style-11')


    '''
    para = doc.add_paragraph('Dito isto, face a publicação da portaria conjunta Nº 47, de 14 de julho de 2022 do Tribunal de Justiça do Rio Grande do Norte,', style='Style-1')
    para.add_run(' em conformidade também ao parágrafo único, do art. 906, CPC para fins de devolução à ré do valor depositado nos autos, e seus acréscimos legais,')
    run = para.add_run(' requerer a liberação da quantia pelo sistema SISCONDJ')
    run.font.bold = True
    para.add_run(',  em favor da SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A., cujos dados da conta seguem abaixo: ')
    '''
    return doc


def pet_004_DevolucaoHPRN1(dados_da_pasta,doc):
    doc.add_paragraph('Consoante se verifica nos autos e da documentação que segue em anexo, houve depósito a título de pagamento de honorários periciais, '+
    em cumprimento à intimação de fls., contudo, o processo foi relacionado para evento de mutirão de perícias, ocasião em que houve o pagamento da prova\
    através de ofício único, restando, portanto, pagamento em duplicidade.', style='Style-2')

    para = doc.add_paragraph('Dito isto, face a publicação da portaria conjunta Nº 47, de 14 de julho de 2022 do Tribunal de Justiça do Rio Grande do Norte,', style='Style-2')
    para.add_run(' em conformidade também ao parágrafo único, do art. 906, CPC para fins de devolução à ré do valor depositado nos autos, e seus acréscimos legais,')
    run = para.add_run(' requerer a liberação da quantia pelo sistema SISCONDJ')
    run.font.bold = True
    para.add_run(',  em favor da SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A., cujos dados da conta seguem abaixo: ')

    doc.add_paragraph('BANCO DO BRASIL S/A', style='Style-1')

    doc.add_paragraph('Titularidade: SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A', style='Style-1')

    doc.add_paragraph('CNPJ/MF: 09.248.608/0001-04', style='Style-1')

    doc.add_paragraph(' Conta corrente nº 644000-2', style='Style-1')

    doc.add_paragraph('Agência: 1912-7.', style='Style-1')

    para = doc.add_paragraph('Necessário esclarecer que a expedição da ordem de pagamento deverá ser nominal à ', style='Style-2')
    run = para.add_run('SEGURADORA LÍDER DOS CONSÓRCIOS DO SEGURO DPVAT S/A')
    para.add_run(', pois foi a empresa que custeou com o depósito como também é a ')
    run = para.add_run('GESTORA')
    run.font.bold = True
    para.add_run('dos ')
    run = para.add_run('Consórcios do Seguro DPVAT nos termos do art. 5º, §3º, da Resolução CNSP de nº 154')
    run.font.bold = True
    para.add_run(', sendo a única e exclusiva beneficiária de reembolso da quantia disponível ao juízo.')

    para = doc.add_paragraph('Reforçando o acima exposto, temos que as regras e os critérios para o DPVAT referentes aos sinistros ocorridos ',style = 'Style-2')
    run = para.add_run('até 31 de dezembro de 2020')
    run.font.bold = True
    para.add_run(' estão estabelecidas, também, na Resolução n.º 399 do CNSP de 29/12/2020.')

    para = doc.add_paragraph('A referida Resolução prevê, no seu artigo 21, a competência da Seguradora Líder:')


    artigoVinteUmResolucao399(doc)

    doc.add_paragraph('Vejamos, agora, o art. 1º da Resolução 400 do CNSP de 29/12/2020:', style = 'Style-2')

    artigoPrimeiroResolucao400(doc)

    doc.add_paragraph('Requer ainda, seja disponibilizado nos autos, extrato que comprove a efetiva transferência conforme estabelecido '+
    'no art. 5 º, Parágrafo Único da Portaria Conjunta Nº 47 do TJRN, possibilitando ao patrono da Ré realizar prestação de contas '+
    'com maior clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.', style = 'Style-2')

    doc.add_paragraph('Nestes Termos,\nPede Deferimento,',style='Style-11')


def artigoPrimeiroResolucao400(doc):
    para = doc.add_paragraph('Art. 1º ', style ='Style-3')
    run = para.add_run('Ratificar que a Seguradora Líder do Consórcio do Seguro DPVAT S.A.')
    run.font.bold = True
    para.add_run(' será a ')
    run = para.add_run('responsável')
    run.font.bold = True
    para.add_run(' pela gestão e operacionalização do seguro')
    run = para.add_run('DPVAT')
    run.font.bold = True
    para.add_run(' referentes, exclusivamente, ')
    run = para.add_run('aos sinistros ocorridos até 31 de dezembro de 2020')
    run.font.bold = True
    run = para.add_run(', inclusive em relação às respectivas ações judiciais posteriormente ajuizadas.')
    run.font.underline = True



def artigoVinteUmResolucao399(doc):
    para = doc.add_paragraph('Art. 21º ', style='Style-3')
    run = para.add_run('A seguradora líder')
    run.font.bold = True
    para.add_run(' do Consórcio DPVAT será ')
    run = para.add_run('responsável')
    run.font.bold = True
    para.add_run(' pela gestão e operacionalização do seguro ')
    run = para.add_run('DPVAT')
    run.font.bold = True
    para.add_run(' referentes, exclusivamente, ')
    run = para.add_run('aos sinistros ocorridos até 31 de dezembro de 2020')
    run.font.bold = True
    para.add_run(' (run-off), ')
    run = para.add_run('inclusive em relação às respectivas ações judiciais posteriormente ajuizadas.')
    run.font.underline =  True

'''
def artigoVinteUmResolucao399(doc):
    para = doc.add_paragraph('Art. 21. ', style='Style-3')
    para.add_run('A seguradora líder ').bold = True
    para.add_run('do Consórcio DPVAT será ')
    para.add_run('responsável ').bold= True
    para.add_run('pela gestão e operacionalização do seguro ')
    para.add_run('DPVAT ').bold = True
    para.add_run('referentes, exclusivamente, ')
    para.add_run('aos sinistros ocorridos até 31 de dezembro de 2020 ').bold = True
    para.add_run('(run-off), ')
    para.add_run('inclusive em relação às respectivas ações judiciais posteriormente ajuizadas.XX').underline = True
'''



def pet_005_DevolucaoHPRN2(dados_da_pasta,doc):

    doc.add_paragraph('Consoante se verifica nos autos e da documentação que segue em anexo, houve depósito a título de pagamento de honorários periciais, '+
    'em cumprimento à intimação de fls., contudo, o processo foi relacionado para evento de mutirão de perícias, ocasião em que houve o pagamento da prova '+
    'através de ofício único, restando, portanto, pagamento em duplicidade.', style='Style-2')

    para = doc.add_paragraph('Dito isto, face a publicação da portaria conjunta Nº 47, de 14 de julho de 2022 do Tribunal de Justiça do Rio Grande do Norte,', style='Style-2')
    para.add_run(' em conformidade também ao parágrafo único, do art. 906, CPC para fins de devolução à ré do valor depositado nos autos, e seus acréscimos legais,')
    run = para.add_run(' requerer a liberação da quantia pelo sistema SISCONDJ')
    run.font.bold = True
    para.add_run(',  em favor da SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A., cujos dados da conta seguem abaixo: ')

    doc.add_paragraph('BANCO DO BRASIL S/A', style='Style-1')

    doc.add_paragraph('Titularidade: SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A', style='Style-1')

    doc.add_paragraph('CNPJ/MF: 09.248.608/0001-04', style='Style-1')

    doc.add_paragraph(' Conta corrente nº 644000-2', style='Style-1')

    doc.add_paragraph('Agência: 1912-7.', style='Style-1')

    para = doc.add_paragraph('Necessário esclarecer que a expedição da ordem de pagamento deverá ser nominal à ', style='Style-2')
    run = para.add_run('SEGURADORA LÍDER DOS CONSÓRCIOS DO SEGURO DPVAT S/A')
    para.add_run(', pois foi a empresa que custeou com o depósito como também é a ')
    run = para.add_run('GESTORA')
    run.font.bold = True
    para.add_run('dos ')
    run = para.add_run('Consórcios do Seguro DPVAT nos termos do art. 5º, §3º, da Resolução CNSP de nº 154')
    run.font.bold = True
    para.add_run(', sendo a única e exclusiva beneficiária de reembolso da quantia disponível ao juízo.')

    para = doc.add_paragraph('Reforçando o acima exposto, temos que as regras e os critérios para o DPVAT referentes aos sinistros ocorridos ',style = 'Style-2')
    run = para.add_run('até 31 de dezembro de 2020')
    run.font.bold = True
    para.add_run(' estão estabelecidas, também, na Resolução n.º 399 do CNSP de 29/12/2020.')

    para = doc.add_paragraph('A referida Resolução prevê, no seu artigo 21, a competência da Seguradora Líder:')

    artigoVinteUmResolucao399(doc)

    doc.add_paragraph('Vejamos, agora, o art. 1º da Resolução 400 do CNSP de 29/12/2020:', style = 'Style-2')

    artigoPrimeiroResolucao400(doc)

    doc.add_paragraph('Requer ainda, seja disponibilizado nos autos, extrato que comprove a efetiva transferência conforme estabelecido '+
    'no art. 5 º, Parágrafo Único da Portaria Conjunta Nº 47 do TJRN, possibilitando ao patrono da Ré realizar prestação de contas '+
    'com maior clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.', style = 'Style-2')

    doc.add_paragraph('Nestes Termos,\nPede Deferimento,',style='Style-11')



def pet_006_DevolucaoHPRN3(dados_da_pasta,doc):

    doc.add_paragraph('Conforme já peticionado nos autos, existem valores a serem devolvidos ao Réu, conforme toda documentação já apresentada e novamente juntada.', style = 'Style-2')
    doc.add_paragraph('Contudo, diante da ausência da parte autora à prova designada, imprescindível para análise do pedido reclamado, o processo foi julgado '+
    'improcedente, decisão esta que já transitou em julgado, merecendo o aludido valor depositado a título de honorários periciais, ser restituído à parte ré.', style = 'Style-2')

    para = doc.add_paragraph('Dito isto, face a publicação da portaria conjunta Nº 47, de 14 de julho de 2022 do Tribunal de Justiça do Rio Grande do Norte,', style='Style-2')
    para.add_run(' em conformidade também ao parágrafo único, do art. 906, CPC para fins de devolução à ré do valor depositado nos autos, e seus acréscimos legais,')
    run = para.add_run(' requerer a liberação da quantia pelo sistema SISCONDJ')
    run.font.bold = True
    para.add_run(',  em favor da SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A., cujos dados da conta seguem abaixo: ')

    doc.add_paragraph('BANCO DO BRASIL S/A', style='Style-1')

    doc.add_paragraph('Titularidade: SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A', style='Style-1')

    doc.add_paragraph('CNPJ/MF: 09.248.608/0001-04', style='Style-1')

    doc.add_paragraph(' Conta corrente nº 644000-2', style='Style-1')

    doc.add_paragraph('Agência: 1912-7.', style='Style-1')

    para = doc.add_paragraph('Necessário esclarecer que a expedição da ordem de pagamento deverá ser nominal à ', style='Style-2')
    run = para.add_run('SEGURADORA LÍDER DOS CONSÓRCIOS DO SEGURO DPVAT S/A')
    para.add_run(', pois foi a empresa que custeou com o depósito como também é a ')
    run = para.add_run('GESTORA')
    run.font.bold = True
    para.add_run('dos ')
    run = para.add_run('Consórcios do Seguro DPVAT nos termos do art. 5º, §3º, da Resolução CNSP de nº 154')
    run.font.bold = True
    para.add_run(', sendo a única e exclusiva beneficiária de reembolso da quantia disponível ao juízo.')

    para = doc.add_paragraph('Reforçando o acima exposto, temos que as regras e os critérios para o DPVAT referentes aos sinistros ocorridos ',style = 'Style-2')
    run = para.add_run('até 31 de dezembro de 2020')
    run.font.bold = True
    para.add_run(' estão estabelecidas, também, na Resolução n.º 399 do CNSP de 29/12/2020.')

    para = doc.add_paragraph('A referida Resolução prevê, no seu artigo 21, a competência da Seguradora Líder:')

    artigoVinteUmResolucao399(doc)

    doc.add_paragraph('Vejamos, agora, o art. 1º da Resolução 400 do CNSP de 29/12/2020:', style = 'Style-2')

    artigoPrimeiroResolucao400(doc)

    doc.add_paragraph('Requer ainda, seja disponibilizado nos autos, extrato que comprove a efetiva transferência conforme estabelecido '+
    'no art. 5 º, Parágrafo Único da Portaria Conjunta Nº 47 do TJRN, possibilitando ao patrono da Ré realizar prestação de contas '+
    'com maior clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.', style = 'Style-2')

    doc.add_paragraph('Nestes Termos,\nPede Deferimento,',style='Style-11')



def montarNomeArquivo(modelo,cod_cliente,pasta):
    return cod_cliente+'_Peticao_'+modelo+'_'+pasta



def pet_007_DevolucaoHPRN4(dados_da_pasta,doc):

    doc.add_paragraph('Em cumprimento à determinação desse d. juízo, a ré procedeu com o pagamento dos honorários periciais.', style = 'Style-2')

    doc.add_paragraph('Contudo, diante da ausência da parte autora à prova designada, imprescindível para análise do pedido reclamado, o processo foi '+
    'julgado extinto sem resolução de mérito, decisão esta que já transitou em julgado, merecendo o aludido valor depositado a título de honorários '+
    'periciais, ser restituído à parte ré.', style = 'Style-2')

    para = doc.add_paragraph('Dito isto, face a publicação da portaria conjunta Nº 47, de 14 de julho de 2022 do Tribunal de Justiça do Rio Grande do Norte,', style='Style-2')
    para.add_run(' em conformidade também ao parágrafo único, do art. 906, CPC para fins de devolução à ré do valor depositado nos autos, e seus acréscimos legais,')
    run = para.add_run(' requerer a liberação da quantia pelo sistema SISCONDJ')
    run.font.bold = True
    para.add_run(',  em favor da SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A., cujos dados da conta seguem abaixo: ')

    doc.add_paragraph('BANCO DO BRASIL S/A', style='Style-1')

    doc.add_paragraph('Titularidade: SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A', style='Style-1')

    doc.add_paragraph('CNPJ/MF: 09.248.608/0001-04', style='Style-1')

    doc.add_paragraph(' Conta corrente nº 644000-2', style='Style-1')

    doc.add_paragraph('Agência: 1912-7.', style='Style-1')

    para = doc.add_paragraph('Necessário esclarecer que a expedição da ordem de pagamento deverá ser nominal à ', style='Style-2')
    run = para.add_run('SEGURADORA LÍDER DOS CONSÓRCIOS DO SEGURO DPVAT S/A')
    para.add_run(', pois foi a empresa que custeou com o depósito como também é a ')
    run = para.add_run('GESTORA')
    run.font.bold = True
    para.add_run('dos ')
    run = para.add_run('Consórcios do Seguro DPVAT nos termos do art. 5º, §3º, da Resolução CNSP de nº 154')
    run.font.bold = True
    para.add_run(', sendo a única e exclusiva beneficiária de reembolso da quantia disponível ao juízo.')

    para = doc.add_paragraph('Reforçando o acima exposto, temos que as regras e os critérios para o DPVAT referentes aos sinistros ocorridos ',style = 'Style-2')
    run = para.add_run('até 31 de dezembro de 2020')
    run.font.bold = True
    para.add_run(' estão estabelecidas, também, na Resolução n.º 399 do CNSP de 29/12/2020.')

    para = doc.add_paragraph('A referida Resolução prevê, no seu artigo 21, a competência da Seguradora Líder:')


    artigoVinteUmResolucao399(doc)

    doc.add_paragraph('Vejamos, agora, o art. 1º da Resolução 400 do CNSP de 29/12/2020:', style = 'Style-2')

    artigoPrimeiroResolucao400(doc)

    doc.add_paragraph('Requer ainda, seja disponibilizado nos autos, extrato que comprove a efetiva transferência conforme estabelecido '+
    'no art. 5 º, Parágrafo Único da Portaria Conjunta Nº 47 do TJRN, possibilitando ao patrono da Ré realizar prestação de contas '+
    'com maior clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.', style = 'Style-2')

    doc.add_paragraph('Nestes Termos,\nPede Deferimento,',style='Style-11')




def pet_008_DevolucaoHPRN5(dados_da_pasta,doc):


    doc.add_paragraph('Em cumprimento à determinação desse d. juízo, a ré procedeu com o pagamento dos honorários periciais.',style = 'Style-2')


    doc.add_paragraph('Contudo, diante da ausência da parte autora à prova designada, imprescindível para análise do pedido reclamado, o processo foi julgado '+
    'extinto sem resolução de mérito, decisão esta que já transitou em julgado, merecendo o aludido valor depositado a título de honorários periciais, ser restituído à parte ré.', style = 'Style-2')

    para = doc.add_paragraph('Dito isto, face a publicação da portaria conjunta Nº 47, de 14 de julho de 2022 do Tribunal de Justiça do Rio Grande do Norte,', style='Style-2')
    para.add_run(' em conformidade também ao parágrafo único, do art. 906, CPC para fins de devolução à ré do valor depositado nos autos, e seus acréscimos legais,')
    run = para.add_run(' requerer a liberação da quantia pelo sistema SISCONDJ')
    run.font.bold = True
    para.add_run(',  em favor da SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A., cujos dados da conta seguem abaixo: ')

    doc.add_paragraph('BANCO DO BRASIL S/A', style='Style-1')

    doc.add_paragraph('Titularidade: SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A', style='Style-1')

    doc.add_paragraph('CNPJ/MF: 09.248.608/0001-04', style='Style-1')

    doc.add_paragraph(' Conta corrente nº 644000-2', style='Style-1')

    doc.add_paragraph('Agência: 1912-7.', style='Style-1')

    para = doc.add_paragraph('Necessário esclarecer que a expedição da ordem de pagamento deverá ser nominal à ', style='Style-2')
    run = para.add_run('SEGURADORA LÍDER DOS CONSÓRCIOS DO SEGURO DPVAT S/A')
    para.add_run(', pois foi a empresa que custeou com o depósito como também é a ')
    run = para.add_run('GESTORA')
    run.font.bold = True
    para.add_run('dos ')
    run = para.add_run('Consórcios do Seguro DPVAT nos termos do art. 5º, §3º, da Resolução CNSP de nº 154')
    run.font.bold = True
    para.add_run(', sendo a única e exclusiva beneficiária de reembolso da quantia disponível ao juízo.')

    para = doc.add_paragraph('Reforçando o acima exposto, temos que as regras e os critérios para o DPVAT referentes aos sinistros ocorridos ',style = 'Style-2')
    run = para.add_run('até 31 de dezembro de 2020')
    run.font.bold = True
    para.add_run(' estão estabelecidas, também, na Resolução n.º 399 do CNSP de 29/12/2020.')

    para = doc.add_paragraph('A referida Resolução prevê, no seu artigo 21, a competência da Seguradora Líder:')

    artigoVinteUmResolucao399(doc)

    doc.add_paragraph('Vejamos, agora, o art. 1º da Resolução 400 do CNSP de 29/12/2020:', style = 'Style-2')

    artigoPrimeiroResolucao400(doc)

    doc.add_paragraph('Requer ainda, seja disponibilizado nos autos, extrato que comprove a efetiva transferência conforme estabelecido '+
    'no art. 5 º, Parágrafo Único da Portaria Conjunta Nº 47 do TJRN, possibilitando ao patrono da Ré realizar prestação de contas '+
    'com maior clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.', style = 'Style-2')

    doc.add_paragraph('Nestes Termos,\nPede Deferimento,',style='Style-11')


def pet_009_DevolucaoHPRN6(dados_da_pasta,doc):
    pet_008_DevolucaoHPRN5(dados_da_pasta,doc)


def pet_010_DevolucaoHPRN7(dados_da_pasta,doc):


    doc.add_paragraph('Em cumprimento à determinação desse d. juízo, a ré procedeu com o pagamento dos honorários periciais.', style = 'Style-2')

    doc.add_paragraph('Contudo, diante da ausência da parte autora à prova designada, imprescindível para análise do pedido reclamado, o processo foi julgado '+
    'improcedente, decisão esta que já transitou em julgado, merecendo o aludido valor depositado a título de honorários periciais, ser restituído à parte ré.', style = 'Style-2')

    para = doc.add_paragraph('Dito isto, face a publicação da portaria conjunta Nº 47, de 14 de julho de 2022 do Tribunal de Justiça do Rio Grande do Norte,', style='Style-2')
    para.add_run(' em conformidade também ao parágrafo único, do art. 906, CPC para fins de devolução à ré do valor depositado nos autos, e seus acréscimos legais,')
    run = para.add_run(' requerer a liberação da quantia pelo sistema SISCONDJ')
    run.font.bold = True
    para.add_run(',  em favor da SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A., cujos dados da conta seguem abaixo: ')

    doc.add_paragraph('BANCO DO BRASIL S/A', style='Style-1')

    doc.add_paragraph('Titularidade: SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A', style='Style-1')

    doc.add_paragraph('CNPJ/MF: 09.248.608/0001-04', style='Style-1')

    doc.add_paragraph(' Conta corrente nº 644000-2', style='Style-1')

    doc.add_paragraph('Agência: 1912-7.', style='Style-1')

    para = doc.add_paragraph('Necessário esclarecer que a expedição da ordem de pagamento deverá ser nominal à ', style='Style-2')
    run = para.add_run('SEGURADORA LÍDER DOS CONSÓRCIOS DO SEGURO DPVAT S/A')
    para.add_run(', pois foi a empresa que custeou com o depósito como também é a ')
    run = para.add_run('GESTORA')
    run.font.bold = True
    para.add_run('dos ')
    run = para.add_run('Consórcios do Seguro DPVAT nos termos do art. 5º, §3º, da Resolução CNSP de nº 154')
    run.font.bold = True
    para.add_run(', sendo a única e exclusiva beneficiária de reembolso da quantia disponível ao juízo.')

    para = doc.add_paragraph('Reforçando o acima exposto, temos que as regras e os critérios para o DPVAT referentes aos sinistros ocorridos ',style = 'Style-2')
    run = para.add_run('até 31 de dezembro de 2020')
    run.font.bold = True
    para.add_run(' estão estabelecidas, também, na Resolução n.º 399 do CNSP de 29/12/2020.')

    para = doc.add_paragraph('A referida Resolução prevê, no seu artigo 21, a competência da Seguradora Líder:')

    artigoVinteUmResolucao399(doc)

    doc.add_paragraph('Vejamos, agora, o art. 1º da Resolução 400 do CNSP de 29/12/2020:', style = 'Style-2')

    artigoPrimeiroResolucao400(doc)

    doc.add_paragraph('Requer ainda, seja disponibilizado nos autos, extrato que comprove a efetiva transferência conforme estabelecido '+
    'no art. 5 º, Parágrafo Único da Portaria Conjunta Nº 47 do TJRN, possibilitando ao patrono da Ré realizar prestação de contas '+
    'com maior clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.', style = 'Style-2')

    doc.add_paragraph('Nestes Termos,\nPede Deferimento,',style='Style-11')



def pet_018_JuntadaDeCustasFinais(dados_da_pasta,doc):
    styles = doc.styles
    p = styles.add_style("Style-6", WD_STYLE_TYPE.PARAGRAPH)
    p.font.name = "Calibri"
    p.font.size = Pt(11)
    #p.font.bold = True
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph('Desta forma, requer a remessa imediata ao juízo a quo, destacando a JUNTADA DAS CUSTAS FINAIS, bem como, pugna-se para que, no juízo de grau mínimo, caso verificado saldo '+
    'remanescente a ser recolhido, seja a demandada intimada em nome do seu causídico abaixo apontado.',style = 'Style-6')

    para = doc.add_paragraph('Por fim, que seja observado exclusivamente o nome ', style = 'Style-6')

    if dados_da_pasta.publicando_sexo=='M':
        para.add_run('do advogado ')
    else:
        para.add_run('da advogada ')
    para.add_run(dados_da_pasta.publicando_nome+' '+dados_da_pasta.publicando_oab).bold = True
    para.add_run(', para efeito de intimações futuras, sob pena de nulidade das mesmas.')

    doc.add_paragraph('Termos em que,\nPede Juntada,',style='Style-11')

    return doc

