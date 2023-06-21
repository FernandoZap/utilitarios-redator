# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Length
import docx

import os
import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches,Cm
from . import funcoes_banco,funcoes_gerais



def pagina_inicial(doc,pasta):
    dados_view = funcoes_banco.f001_sql(pasta,1)
    dados_pasta = funcoes_banco.f002_sql(pasta)
    nr_processo = fun_nr_processo(dados_view['nr_processo'],dados_view['nr_processoSE'],dados_view['uf'])

    styles = doc.styles
    p4 = styles.add_style("style4", WD_STYLE_TYPE.PARAGRAPH)
    p4.font.name = "Calibri"
    p4.font.size = Pt(11)
    #p4.font.color.rgb=RGBColor(79, 129, 189)
    p4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #p4.paragraph_format.first_line_indent = Inches(0.5)
    #p4.paragraph_format.left_indent = Inches(0.5)



    styles2 = doc.styles
    p2 = styles.add_style("styles2", WD_STYLE_TYPE.PARAGRAPH)
    p2.font.name = "Calibri"
    p2.font.bold = True
    p2.font.size = Pt(11)
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


    font = doc.styles['Normal'].font
    font.name='Calibri'
    font.size=Pt(11)


    # paragrafo acima da logo jbaa
    paragraph1 = doc.add_paragraph()
    paragraph1_para = paragraph1.add_run(dados_view['cod_cliente'] +' - C1/ '+ dados_view['pasta']+ '/ '+dados_view['secao'])
    paragraph1.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    paragraph1_para.font.size = Pt(11)

    # construcao da primeira pagina
    #imagem= doc.add_picture('C:/Users/Fernando/saj vps/img/logo_jbaa.png', width=Inches(1.25))
    imagem= doc.add_picture(os.environ.get('DIR_IMAGENS')+'/logo_jbaa.png', width=Inches(1.45))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph

    # paragrafo abaixo da logo jbaa
    paragraph2 = doc.add_paragraph()
    string_vara = (dados_pasta['preposicao'] + ' ' + dados_pasta['num_orgao'] + ' '+ dados_pasta['orgao']).upper()
    paragraph2 = paragraph2.add_run('\nEXMO.SR.DR.JUIZ DE DIREITO ' + string_vara +' DA COMARCA DE ' + dados_pasta['comarca']  +'/'+dados_pasta['estado']+'\n')
    paragraph2.bold = True

    # nº processo
    paragraph3 = doc.add_paragraph()
    paragraph3.add_run('Processo: '+nr_processo+'\n').bold = True
    #paragraph3_para.alignment=WD_ALIGN_PARAGRAPH.LEFT
    #paragraph3.bold = True

    paragraph4 = doc.add_paragraph('',style='style4')
    paragraph4.add_run(dados_pasta['reu']).bold = True

    paragraph4.add_run(', já devidamente qualificado nos autos do processo em epígrafe, por meio de seus advogados que esta subscreve,\
        vem à presença de V. Excelência, nos autos da AÇÃO DE COBRANÇA DE SEGURO DPVAT promovida por '+ dados_pasta['autor']+', opor \
        EMBARGOS DE DECLARAÇÃO, conforme passa a expor: ')


    #paragraph4.paragraph_format.first_line_indent = Inches(0.5)
    #paragraph4.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY


    # footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    para = footer_para.add_run(\
        ' \tRua São José, 90, 8º andar, Centro, Rio de Janeiro/RJ - CEP: 20010-020 \
        \n\twww.joaobarbosaadvass.com.br'\
    )
    para.font.size = Pt(9)
    para.font.name = 'Calibri'
    #doc.add_page_break()





def pagina_conclusao(doc,pasta):
    doc.add_paragraph('CONCLUSÃO')

    doc.add_paragraph('São essas as razões pelas quais a embargante confia, espera e requer sejam acolhidos e providos os presentes Embargos Declaratórios, '+
    'enfrentado o ponto OMISSO, conferido efeitos integrativos para o fim de prover integralmente, tudo por ser medida de direito e irretorquível JUSTIÇA!')



def pagina_de_encerramento(doc,pasta):
    dados_view = funcoes_banco.f001_sql(pasta,1)
    dados_pasta = funcoes_banco.f002_sql(pasta)

    paragraph4=doc.add_paragraph('Neste Termos')
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph4.paragraph_format.space_after = Pt(1)
    paragraph4=doc.add_paragraph('Pede Deferimento')
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    local = funcoes_gerais.local_e_data(dados_pasta['comarca'])
    paragraph4=doc.add_paragraph(local)
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER

    paragraph4=doc.add_paragraph()
    paragraph4_para = paragraph4.add_run('\nJOÃO BARBOSA')
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph4.paragraph_format.space_before = Pt(1)
    paragraph4.paragraph_format.space_after = Pt(0)
    paragraph4_para.bold = True

    paragraph4=doc.add_paragraph()
    paragraph4_para = paragraph4.add_run(dados_view['publicando_nome']+' - '+dados_view['publicando_oab'])
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph4_para.bold = True

    paragraph4=doc.add_paragraph()
    paragraph4_para = paragraph4.add_run('\n'+dados_pasta['advConveniado'])
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph4.paragraph_format.space_before = Pt(1)
    paragraph4.paragraph_format.space_after = Pt(0)
    paragraph4_para.bold = True

    paragraph4=doc.add_paragraph()
    paragraph4_para = paragraph4.add_run(dados_pasta['oabConveniado'])
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph4.paragraph_format.space_before = Pt(1)
    paragraph4_para.bold = True


def pagina_inicial_ultrapetita(doc,pasta):
    pass

def sinteseDosFatos(doc):
    paragraph2 = doc.add_paragraph()
    paragraph2 = paragraph2.add_run('DA SÍNTESE DOS FATOS E DA OMISSÃO NA DECISÃO PROFERIDA')
    paragraph2.bold = True
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph


    paragraph3 = doc.add_paragraph()
    paragraph3 = paragraph3.add_run('Sem adentrar ao mérito do decisum, informa a V.Exa. que constou na parte dispositiva '+
    'desta o seguinte:')

    paragraph4 = doc.add_paragraph()
    paragraph4 = paragraph4.add_run('" COLACIONAR SENTENÇA  "')
    paragraph4.bold = True
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph


    paragraph5 = doc.add_paragraph()
    paragraph5 = paragraph5.add_run('Com a mais respeitosa vênia, assim o fazendo afigura-se a v. decisão omissa em pontos essenciais, '+
    'justificando o cabimento dos presentes Embargos de Declaração, a fim de que essa V. Exa. decida-os e  confira os efeitos integrativos '+
    'ao respeitpavel decisum.')
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph


    paragraph6 = doc.add_paragraph()
    paragraph6 = paragraph6.add_run('Verifica-se grava OMISSÃO, que devem ser supridas ou sanadas por meio dos presentes embargos, '+
    'sendo certo que o recurso não objetiva rediscutir a matéria, mas afastar os vícios constatados no julgado')
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.add_paragraph

    #doc.add_page_break()


def pagina_inicial_peticoes_desarquivamento(doc,dados_da_pasta,tipo,informs_da_pasta):
    dados_view = funcoes_banco.f001_sql(dados_da_pasta.pasta,1)
    dados_pasta = funcoes_banco.f002_sql(dados_da_pasta.pasta)

    nr_processo=fun_nr_processo(dados_view['nr_processo'],dados_view['nr_processoSE'],dados_view['uf'])

    styles = doc.styles
    p4 = styles.add_style("style4", WD_STYLE_TYPE.PARAGRAPH)
    p4.font.name = "Calibri"
    p4.font.size = Pt(11)
    p4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if tipo==9:
        p4.paragraph_format.first_line_indent = Inches(0.5)


    # paragrafo acima da logo jbaa
    paragraph1 = doc.add_paragraph()
    paragraph1_para = paragraph1.add_run(dados_view['cod_cliente'] +' - C1/ '+ dados_view['pasta']+ '/ '+dados_view['secao'])
    paragraph1.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    paragraph1_para.font.size = Pt(11)

    # construcao da primeira pagina
    #imagem= doc.add_picture('C:/Users/Fernando/saj vps/img/logo_jbaa.png', width=Inches(1.25))
    imagem= doc.add_picture(os.environ.get('DIR_IMAGENS')+'/logo_jbaa.png', width=Inches(1.45))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph

    # paragrafo abaixo da logo jbaa
    paragraph2 = doc.add_paragraph()
    string_vara = (dados_pasta['preposicao'] + ' ' + dados_pasta['num_orgao'] + ' '+ dados_pasta['orgao']).upper()
    paragraph2 = paragraph2.add_run('\nEXMO.SR.DR.JUIZ DE DIREITO ' + string_vara +' DA COMARCA DE ' + dados_pasta['comarca']  +'/'+dados_pasta['estado']+'\n')
    paragraph2.bold = True

    # nº processo
    paragraph3 = doc.add_paragraph()
    paragraph3.add_run('Processo: '+nr_processo+'\n').bold = True


    para = doc.add_paragraph('',style='style4')
    para.add_run(dados_pasta['reu']).bold = True
    para.add_run(', previamente qualificada nos autos do processo em epígrafe, neste ato, representada por seus advogados que esta subscrevem, nos autos da ')
    para.add_run('AÇÃO DE COBRANÇA DE SEGURO DPVAT').bold = True
    #run.font.bold = True
    para.add_run(', que lhe promove ')
    para.add_run(dados_pasta['autor']).bold = True
    #run.font.bold = True
    #run.font.underline = True

    if tipo==1:
        para.add_run(', em trâmite perante este Douto Juízo, vem respeitosamente, à presença de V. Exa., requerer o ')
        para.add_run('DESARQUIVAMENTO').bold = True
        #run.font.bold = True
        para.add_run(', a fim de viabilizar a DEVOLUÇÃO DOS HONORÁRIOS PERICIAIS PAGOS EM DUPLICIDADE (depósito judicial e oficio único de pagamento).').bold = True
        #run.font.bold = True
    elif tipo==2:
        para.add_run(', em trâmite perante este Douto Juízo e Respectivo Cartório, vem, mui respeitosamente, à presença de V. Exa., pugnar pelo DESARQUIVAMENTO DO AUTOS, para após informar e requerer o que segue:')
    elif tipo==3:
        para.add_run(', em trâmite perante este Douto Juízo, vem respeitosamente, à presença de V. Exa., requerer que seja determinada a juntada de ')
        run = para.add_run('RECIBO DE PAGAMENTO E OFÍCIO')
        run.font.bold = True
        para.add_run(' em anexo, com fito de ')
        run = para.add_run('comprovar o pagamento dos honorários do perito nomeado pelo Juízo.')
        run.font.bold = True
        doc.add_paragraph('Termo em que,\nPede Deferimento.', style='style4')
    elif tipo==4:
        para.add_run(', em trâmite perante este Douto Juízo e Respectivo Cartório, vem, mui respeitosamente, à presença de V. Exa., ')
        run = para.add_run('requerer o DESARQUIVAMENTO, a fim de viabilizar a DEVOLUÇÃO DOS HONORÁRIOS PERICIAIS PAGOS EM DUPLICIDADE (depósito judicial e ofício único de pagamento).')
        run.font.bold = True
    elif tipo==5:
        para.add_run(', em trâmite perante este Douto Juízo e Respectivo Cartório, vem, mui respeitosamente, à presença de V. Exa., ')
        run = para.add_run('requerer a DEVOLUÇÃO DOS HONORÁRIOS PERICIAIS PAGOS EM DUPLICIDADE (depósito judicial e ofício único de pagamento).')
        run.font.bold = True
    elif tipo==6:
        para.add_run(', em trâmite perante este Douto Juízo e Respectivo Cartório, vem, mui respeitosamente, à presença de V. Exa., inicialmente pugnar pelo DESARQUIVAMENTO DOS AUTOS, para informar e requerer o que segue:')
    elif tipo==7:
        para.add_run(', em trâmite perante este Douto Juízo e Respectivo Cartório, vem, mui respeitosamente, à presença de V. Exa., informar e requerer o que segue:')
    elif tipo==8:
        para.add_run(', em trâmite perante este Douto Juízo, vem, respeitosamente, à presença de V. Exa.,')
        run = para.add_run(' requerer a juntada do Comprovante de Pagamento da liquidação, no valor de R$ ')
        run.font.bold = True
        valor = informs_da_pasta['valorDoPagamento']
        run = para.add_run(valor+' ('+(funcoes_gerais.valor_por_extenso(valor)).upper()+').')
        run.font.bold = True
        #TRÊS MIL E SEISCENTOS E OITENTA E CINCO REAIS E SETENTA E CINCO CENTAVOS)')
    elif tipo==9:
        para.add_run(', em trâmite perante este Douto Juízo, vem, respeitosamente, à presença de V. Exa.,')
        run = para.add_run(' requerer a juntada da inclusa guia de recolhimento de custas finais.')
        run.font.bold = True
    elif tipo==10:
        para.add_run(', em trâmite perante este Douto Juízo e Respectivo Cartório, vem, mui respeitosamente, à presença de V. Exa., informar para ao final requerer o que segue:')


    # footer
    '''
    section = doc.sections[0]
    footer = section.footer
    footer_para = xfooter.paragraphs[0]
    para = footer_para.add_run(\
        ' \tRua São José, 90, 8º andar, Centro, Rio de Janeiro/RJ - CEP: 20010-020 \
        \n\twww.joaobarbosaadvass.com.br'\
    )
    para.font.size = Pt(9)
    para.font.name = 'Calibri'
    '''


def pagina_inicial_peticoes_devolucao(doc,dados_da_pasta):
    dados_view = funcoes_banco.f001_sql(dados_da_pasta.pasta,1)
    dados_pasta = funcoes_banco.f002_sql(dados_da_pasta.pasta)


    nr_processo=fun_nr_processo(dados_view['nr_processo'],dados_view['nr_processoSE'],dados_view['uf'])
    styles = doc.styles
    p4 = styles.add_style("style4", WD_STYLE_TYPE.PARAGRAPH)
    p4.font.name = "Calibri"
    p4.font.size = Pt(11)
    p4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    font = doc.styles['Normal'].font
    font.name='Calibri'
    font.size=Pt(11)


    # margin doc
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # paragrafo acima da logo jbaa
    paragraph1 = doc.add_paragraph()
    paragraph1_para = paragraph1.add_run(dados_view['cod_cliente'] +' - C1/ '+ dados_view['pasta']+ '/ '+dados_view['secao'])
    paragraph1.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    paragraph1_para.font.size = Pt(11)

    # construcao da primeira pagina
    #imagem= doc.add_picture('C:/Users/Fernando/saj vps/img/logo_jbaa.png', width=Inches(1.25))
    imagem= doc.add_picture(os.environ.get('DIR_IMAGENS')+'/logo_jbaa.png', width=Inches(1.45))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph

    # paragrafo abaixo da logo jbaa
    paragraph2 = doc.add_paragraph()
    string_vara = (dados_pasta['preposicao'] + ' ' + dados_pasta['num_orgao'] + ' '+ dados_pasta['orgao']).upper()
    paragraph2 = paragraph2.add_run('\nEXMO.SR.DR.JUIZ DE DIREITO ' + string_vara +' DA COMARCA DE ' + dados_pasta['comarca']  +'/'+dados_pasta['estado']+'\n')
    paragraph2.bold = True

    # nº processo
    paragraph3 = doc.add_paragraph()
    paragraph3.add_run('Processo: '+nr_processo+'\n').bold = True

    paragraph4 = doc.add_paragraph('',style='style4')
    paragraph4.add_run(dados_pasta['reu']).bold = True
    paragraph4.add_run(', previamente qualificada nos autos do processo em epígrafe, neste ato, representada por seus advogados que esta subscrevem, nos autos da ')
    run = paragraph4.add_run('AÇÃO DE COBRANÇA DE SEGURO DPVAT')
    run.font.bold = True
    paragraph4.add_run(', que lhe promove ')
    para = paragraph4.add_run(dados_pasta['autor'])
    para.font.bold = True
    para.font.underline = True

    paragraph4.add_run(', em trâmite perante este Douto Juízo e Respectivo Cartório, vem, mui respeitosamente, '
    'à presença de V. Exa., informar para ao final requerer o que segue:')


    # footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    para = footer_para.add_run(\
        ' \tRua São José, 90, 8º andar, Centro, Rio de Janeiro/RJ - CEP: 20010-020 \
        \n\twww.joaobarbosaadvass.com.br'\
    )
    para.font.size = Pt(9)
    para.font.name = 'Calibri'



def pagina_de_encerramento_peticoes(doc,pasta):
    dados_view = funcoes_banco.f001_sql(pasta,1)
    dados_pasta = funcoes_banco.f002_sql(pasta)
    '''
    paragraph4=doc.add_paragraph('Neste Termos')
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph4.paragraph_format.space_after = Pt(1)
    paragraph4=doc.add_paragraph('Pede Deferimento')
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    '''
 
    local = funcoes_gerais.local_e_data(dados_pasta['comarca'])
    paragraph4=doc.add_paragraph(local)
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER

    paragraph4=doc.add_paragraph()
    paragraph4_para = paragraph4.add_run('\nJOÃO BARBOSA')
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph4.paragraph_format.space_before = Pt(1)
    paragraph4.paragraph_format.space_after = Pt(0)
    paragraph4_para.bold = True

    paragraph4=doc.add_paragraph()
    #paragraph4_para = paragraph4.add_run(dados_view['publicando_nome']+' - OAB/'+dados_view['uf'])
    paragraph4_para = paragraph4.add_run(dados_view['oabJB'])
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph4_para.bold = True

    paragraph4=doc.add_paragraph()
    paragraph4_para = paragraph4.add_run('\n'+dados_view['conveniado_nome'])
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph4.paragraph_format.space_before = Pt(1)
    paragraph4.paragraph_format.space_after = Pt(0)
    paragraph4_para.bold = True

    paragraph4=doc.add_paragraph()
    paragraph4_para = paragraph4.add_run(dados_view['conveniado_oab'])
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph4.paragraph_format.space_before = Pt(1)
    paragraph4_para.bold = True




def fun_nr_processo(numproc,numprocSE,uf):
    nr_processo=numproc

    if uf=='SE':
        if numprocSE!='':
            nr_processo=numprocSE
    return nr_processo









